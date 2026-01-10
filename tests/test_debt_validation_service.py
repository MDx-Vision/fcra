"""
Unit Tests for Debt Validation Service.
Tests for debt validation letter generation, PII extraction, and auto-send functionality.
Covers PDF/Word document generation, client data extraction, and edge cases.
"""
import pytest
from datetime import date, datetime, timedelta
from unittest.mock import patch, MagicMock, PropertyMock
import sys
import os
import tempfile

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from services.debt_validation_service import (
    get_client_pii,
    generate_debt_validation_letter,
    generate_validation_letters,
    generate_validation_letter_single,
    auto_generate_validation_letters_from_analysis,
    get_common_collection_agencies,
    DebtValidationPDF,
    _add_validation_letter_to_pdf,
    _create_validation_letter_docx,
    _add_validation_letter_to_docx,
    COLLECTION_AGENCY_TEMPLATE,
)


# =============================================================================
# Tests for get_client_pii() - Dictionary Input
# =============================================================================

class TestGetClientPiiDict:
    """Test PII extraction from dictionary input."""

    def test_get_client_pii_full_dict(self):
        """Test extracting PII from a complete dictionary."""
        client_dict = {
            "name": "John Smith",
            "first_name": "John",
            "last_name": "Smith",
            "address_street": "123 Main St",
            "address_city": "Los Angeles",
            "address_state": "CA",
            "address_zip": "90001",
            "ssn_last_four": "1234",
            "date_of_birth": date(1985, 6, 15),
            "email": "john@example.com",
            "phone": "555-123-4567",
        }
        pii = get_client_pii(client_dict)

        assert pii["name"] == "John Smith"
        assert pii["first_name"] == "John"
        assert pii["last_name"] == "Smith"
        assert pii["address_street"] == "123 Main St"
        assert pii["address_city"] == "Los Angeles"
        assert pii["address_state"] == "CA"
        assert pii["address_zip"] == "90001"
        assert pii["ssn_last_four"] == "1234"
        assert pii["date_of_birth"] == date(1985, 6, 15)
        assert pii["email"] == "john@example.com"
        assert pii["phone"] == "555-123-4567"

    def test_get_client_pii_dict_name_from_first_last(self):
        """Test name is constructed from first_name and last_name when name is empty."""
        client_dict = {
            "name": "",
            "first_name": "Jane",
            "last_name": "Doe",
        }
        pii = get_client_pii(client_dict)
        assert pii["name"] == "Jane Doe"

    def test_get_client_pii_dict_name_takes_priority(self):
        """Test full name takes priority over first/last name."""
        client_dict = {
            "name": "Robert Johnson Jr.",
            "first_name": "Robert",
            "last_name": "Johnson",
        }
        pii = get_client_pii(client_dict)
        assert pii["name"] == "Robert Johnson Jr."

    def test_get_client_pii_dict_missing_fields(self):
        """Test default values for missing fields."""
        client_dict = {}
        pii = get_client_pii(client_dict)

        assert pii["name"] == ""
        assert pii["first_name"] == ""
        assert pii["last_name"] == ""
        assert pii["address_street"] == ""
        assert pii["ssn_last_four"] == "XXXX"
        assert pii["date_of_birth"] is None
        assert pii["email"] == ""
        assert pii["phone"] == ""

    def test_get_client_pii_dict_partial_data(self):
        """Test with partial client data."""
        client_dict = {
            "first_name": "Alice",
            "last_name": "Wonder",
            "email": "alice@example.com",
        }
        pii = get_client_pii(client_dict)

        assert pii["name"] == "Alice Wonder"
        assert pii["email"] == "alice@example.com"
        assert pii["address_street"] == ""
        assert pii["ssn_last_four"] == "XXXX"


# =============================================================================
# Tests for get_client_pii() - Object Input
# =============================================================================

class TestGetClientPiiObject:
    """Test PII extraction from client object."""

    def test_get_client_pii_full_object(self):
        """Test extracting PII from a complete client object."""
        client = MagicMock()
        client.name = "Michael Brown"
        client.first_name = "Michael"
        client.last_name = "Brown"
        client.address_street = "456 Oak Ave"
        client.address_city = "Chicago"
        client.address_state = "IL"
        client.address_zip = "60601"
        client.ssn_last_four = "5678"
        client.date_of_birth = date(1990, 12, 25)
        client.email = "michael@example.com"
        client.phone = "555-987-6543"

        pii = get_client_pii(client)

        assert pii["name"] == "Michael Brown"
        assert pii["first_name"] == "Michael"
        assert pii["last_name"] == "Brown"
        assert pii["address_street"] == "456 Oak Ave"
        assert pii["address_city"] == "Chicago"
        assert pii["address_state"] == "IL"
        assert pii["address_zip"] == "60601"
        assert pii["ssn_last_four"] == "5678"
        assert pii["date_of_birth"] == date(1990, 12, 25)
        assert pii["email"] == "michael@example.com"
        assert pii["phone"] == "555-987-6543"

    def test_get_client_pii_object_name_from_first_last(self):
        """Test name is constructed from first/last when name is empty."""
        client = MagicMock()
        client.name = ""
        client.first_name = "Sarah"
        client.last_name = "Connor"
        client.address_street = None
        client.address_city = None
        client.address_state = None
        client.address_zip = None
        client.ssn_last_four = None
        client.date_of_birth = None
        client.email = None
        client.phone = None

        pii = get_client_pii(client)
        assert pii["name"] == "Sarah Connor"

    def test_get_client_pii_object_none_values(self):
        """Test default values when object has None values."""
        client = MagicMock()
        client.name = None
        client.first_name = None
        client.last_name = None
        client.address_street = None
        client.address_city = None
        client.address_state = None
        client.address_zip = None
        client.ssn_last_four = None
        client.date_of_birth = None
        client.email = None
        client.phone = None

        pii = get_client_pii(client)

        assert pii["name"] == ""
        assert pii["first_name"] == ""
        assert pii["last_name"] == ""
        assert pii["ssn_last_four"] == "XXXX"
        assert pii["date_of_birth"] is None


# =============================================================================
# Tests for generate_debt_validation_letter()
# =============================================================================

class TestGenerateDebtValidationLetter:
    """Test debt validation letter content generation."""

    def test_generate_letter_basic(self):
        """Test generating a basic debt validation letter."""
        client = {
            "name": "John Test",
            "address_street": "100 Test Lane",
            "address_city": "Test City",
            "address_state": "TX",
            "address_zip": "75001",
            "ssn_last_four": "9999",
            "date_of_birth": date(1980, 5, 10),
        }
        collection_info = {
            "creditor_name": "ABC Collections",
            "creditor_address": "200 Collection Blvd",
            "creditor_city_state_zip": "Dallas, TX 75201",
            "original_creditor": "Original Bank",
            "account_number": "1234567890",
            "balance": 5000.50,
        }

        letter = generate_debt_validation_letter(client, collection_info)

        assert letter["client_name"] == "John Test"
        assert letter["client_address"] == "100 Test Lane"
        assert letter["client_city_state_zip"] == "Test City, TX 75001"
        assert letter["creditor_name"] == "ABC Collections"
        assert letter["creditor_address"] == "200 Collection Blvd"
        assert letter["original_creditor"] == "Original Bank"
        assert letter["account_number"] == "XXXX7890"  # Masked
        assert letter["balance"] == "$5,000.50"
        assert letter["ssn_last4"] == "9999"

    def test_generate_letter_account_masking_short_number(self):
        """Test account number masking for short account numbers."""
        client = {"name": "Test User"}
        collection_info = {
            "account_number": "123",  # Less than 4 digits
            "balance": 100,
        }

        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["account_number"] == "123"  # Not masked

    def test_generate_letter_account_masking_exactly_4(self):
        """Test account number masking for exactly 4 digit account."""
        client = {"name": "Test User"}
        collection_info = {
            "account_number": "1234",
            "balance": 100,
        }

        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["account_number"] == "1234"  # Not masked (exactly 4)

    def test_generate_letter_account_masking_long_number(self):
        """Test account number masking for long account numbers."""
        client = {"name": "Test User"}
        collection_info = {
            "account_number": "9876543210123456",
            "balance": 100,
        }

        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["account_number"] == "XXXX3456"

    def test_generate_letter_balance_formatting_integer(self):
        """Test balance formatting for integer values."""
        client = {"name": "Test User"}
        collection_info = {
            "account_number": "12345",
            "balance": 1500,
        }

        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["balance"] == "$1,500.00"

    def test_generate_letter_balance_formatting_float(self):
        """Test balance formatting for float values."""
        client = {"name": "Test User"}
        collection_info = {
            "account_number": "12345",
            "balance": 12345.67,
        }

        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["balance"] == "$12,345.67"

    def test_generate_letter_balance_formatting_string(self):
        """Test balance formatting when balance is already a string."""
        client = {"name": "Test User"}
        collection_info = {
            "account_number": "12345",
            "balance": "$1,000.00",  # Already formatted
        }

        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["balance"] == "$1,000.00"

    def test_generate_letter_missing_address(self):
        """Test letter generation with missing client address."""
        client = {
            "name": "No Address User",
            "address_street": "",
            "address_city": "",
            "address_state": "",
            "address_zip": "",
        }
        collection_info = {
            "creditor_name": "Collector",
            "account_number": "12345",
            "balance": 500,
        }

        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["client_address"] == "[CLIENT ADDRESS]"
        assert letter["client_city_state_zip"] == "[CITY, STATE ZIP]"

    def test_generate_letter_missing_collection_info(self):
        """Test letter generation with minimal collection info."""
        client = {"name": "Test User"}
        collection_info = {}

        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["creditor_name"] == "[COLLECTION AGENCY]"
        assert letter["creditor_address"] == "[ADDRESS]"
        assert letter["original_creditor"] == "[ORIGINAL CREDITOR]"
        # Account number gets masked since "[ACCOUNT NUMBER]" is > 4 chars
        assert "XXXX" in letter["account_number"]
        assert letter["balance"] == "[AMOUNT]"

    def test_generate_letter_date_of_birth_datetime(self):
        """Test formatting date of birth as datetime object."""
        client = {
            "name": "Test User",
            "date_of_birth": datetime(1985, 6, 15, 10, 30),
        }
        collection_info = {"account_number": "12345", "balance": 100}

        # The function uses the date, letter content doesn't include DOB directly
        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["client_name"] == "Test User"

    def test_generate_letter_date_of_birth_string(self):
        """Test formatting date of birth as string."""
        client = {
            "name": "Test User",
            "date_of_birth": "June 15, 1985",
        }
        collection_info = {"account_number": "12345", "balance": 100}

        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["client_name"] == "Test User"

    def test_generate_letter_reference_number(self):
        """Test letter generation with reference number."""
        client = {"name": "Test User"}
        collection_info = {
            "account_number": "12345",
            "balance": 100,
            "reference_number": "REF-123456",
        }

        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["reference_number"] == "REF-123456"


# =============================================================================
# Tests for DebtValidationPDF Class
# =============================================================================

class TestDebtValidationPDF:
    """Test custom PDF class for debt validation letters."""

    def test_pdf_creation(self):
        """Test PDF object creation."""
        pdf = DebtValidationPDF()
        assert pdf is not None
        # Check auto page break is set
        assert pdf.auto_page_break is True

    def test_pdf_header_empty(self):
        """Test that header method does nothing (empty header)."""
        pdf = DebtValidationPDF()
        pdf.add_page()
        # Header should not raise any errors
        pdf.header()

    def test_pdf_footer(self):
        """Test that footer adds page number."""
        pdf = DebtValidationPDF()
        pdf.add_page()
        # Footer should not raise errors
        pdf.footer()
        # After footer is called, we're still on page 1 or more
        assert pdf.page_no() >= 1


# =============================================================================
# Tests for _add_validation_letter_to_pdf()
# =============================================================================

class TestAddValidationLetterToPdf:
    """Test adding validation letter content to PDF."""

    def test_add_letter_to_pdf_basic(self):
        """Test adding a basic letter to PDF."""
        pdf = DebtValidationPDF()
        letter_content = {
            "date": "January 15, 2024",
            "client_name": "John Doe",
            "client_address": "123 Main St",
            "client_city_state_zip": "Test City, CA 90001",
            "creditor_name": "ABC Collections",
            "creditor_address": "456 Collection Ave",
            "creditor_city_state_zip": "Collector City, TX 75001",
            "original_creditor": "Original Bank",
            "account_number": "XXXX1234",
            "balance": "$5,000.00",
            "ssn_last4": "1234",
            "reference_number": "",
        }

        _add_validation_letter_to_pdf(pdf, letter_content)

        # PDF should have at least one page (letters may span 2 pages due to content)
        assert pdf.page_no() >= 1

    def test_add_letter_to_pdf_with_reference(self):
        """Test adding letter with reference number."""
        pdf = DebtValidationPDF()
        letter_content = {
            "date": "January 15, 2024",
            "client_name": "John Doe",
            "client_address": "123 Main St",
            "client_city_state_zip": "Test City, CA 90001",
            "creditor_name": "ABC Collections",
            "creditor_address": "456 Collection Ave",
            "creditor_city_state_zip": "Collector City, TX 75001",
            "original_creditor": "Original Bank",
            "account_number": "XXXX1234",
            "balance": "$5,000.00",
            "ssn_last4": "1234",
            "reference_number": "REF-999",
        }

        _add_validation_letter_to_pdf(pdf, letter_content)
        # Letter with reference should have at least 1 page
        assert pdf.page_no() >= 1

    def test_add_multiple_letters_to_pdf(self):
        """Test adding multiple letters to same PDF."""
        pdf = DebtValidationPDF()
        letter_content = {
            "date": "January 15, 2024",
            "client_name": "John Doe",
            "client_address": "123 Main St",
            "client_city_state_zip": "Test City, CA 90001",
            "creditor_name": "ABC Collections",
            "creditor_address": "456 Collection Ave",
            "creditor_city_state_zip": "Collector City, TX 75001",
            "original_creditor": "Original Bank",
            "account_number": "XXXX1234",
            "balance": "$5,000.00",
            "ssn_last4": "1234",
            "reference_number": "",
        }

        initial_pages = pdf.page_no()
        _add_validation_letter_to_pdf(pdf, letter_content)
        first_letter_pages = pdf.page_no()
        _add_validation_letter_to_pdf(pdf, letter_content)
        final_pages = pdf.page_no()

        # Second letter should add more pages
        assert final_pages > first_letter_pages


# =============================================================================
# Tests for _create_validation_letter_docx()
# =============================================================================

class TestCreateValidationLetterDocx:
    """Test Word document content creation."""

    def test_create_docx_content_basic(self):
        """Test creating basic document content."""
        letter_content = {
            "date": "January 15, 2024",
            "client_name": "Jane Doe",
            "client_address": "789 Oak St",
            "client_city_state_zip": "Oak City, NY 10001",
            "creditor_name": "XYZ Collections",
            "creditor_address": "321 Debt Lane",
            "creditor_city_state_zip": "Debt City, FL 33101",
            "original_creditor": "Big Bank",
            "account_number": "XXXX5678",
            "balance": "$10,000.00",
            "ssn_last4": "5678",
            "reference_number": "",
        }

        content_items = _create_validation_letter_docx(letter_content)

        assert isinstance(content_items, list)
        assert len(content_items) > 0

        # Check that client name is in content
        names_found = [item for item in content_items
                      if item.get("type") == "paragraph" and item.get("text") == "Jane Doe"]
        assert len(names_found) > 0

    def test_create_docx_content_with_reference(self):
        """Test creating document content with reference number."""
        letter_content = {
            "date": "January 15, 2024",
            "client_name": "Jane Doe",
            "client_address": "789 Oak St",
            "client_city_state_zip": "Oak City, NY 10001",
            "creditor_name": "XYZ Collections",
            "creditor_address": "321 Debt Lane",
            "creditor_city_state_zip": "Debt City, FL 33101",
            "original_creditor": "Big Bank",
            "account_number": "XXXX5678",
            "balance": "$10,000.00",
            "ssn_last4": "5678",
            "reference_number": "REF-12345",
        }

        content_items = _create_validation_letter_docx(letter_content)

        # Check that reference number is included
        ref_found = [item for item in content_items
                    if item.get("type") == "paragraph" and "REF-12345" in item.get("text", "")]
        assert len(ref_found) > 0

    def test_create_docx_content_structure(self):
        """Test document content structure has required sections."""
        letter_content = {
            "date": "January 15, 2024",
            "client_name": "Test Client",
            "client_address": "Test Address",
            "client_city_state_zip": "Test City, ST 00000",
            "creditor_name": "Test Collector",
            "creditor_address": "Collector Address",
            "creditor_city_state_zip": "Collector City, ST 11111",
            "original_creditor": "Original",
            "account_number": "XXXX0000",
            "balance": "$0.00",
            "ssn_last4": "0000",
            "reference_number": "",
        }

        content_items = _create_validation_letter_docx(letter_content)

        # Check for key sections
        texts = [item.get("text", "") for item in content_items if item.get("type") == "paragraph"]
        text_joined = " ".join(texts)

        assert "NOTICE OF DISPUTE AND DEMAND FOR DEBT VALIDATION" in text_joined
        assert "PURSUANT TO 15 U.S.C. 1692g (FDCPA)" in text_joined
        assert "REQUIRED DOCUMENTATION FOR PROPER VALIDATION:" in text_joined
        assert "ADDITIONAL DEMANDS:" in text_joined
        assert "LEGAL NOTICE:" in text_joined


# =============================================================================
# Tests for generate_validation_letters() - Mocked Database
# =============================================================================

class TestGenerateValidationLetters:
    """Test validation letters generation with mocked database."""

    @patch('services.debt_validation_service.get_db')
    @patch('services.debt_validation_service.os.makedirs')
    def test_generate_letters_client_not_found(self, mock_makedirs, mock_get_db):
        """Test error when client is not found."""
        mock_db = MagicMock()
        mock_db.query().filter().first.return_value = None
        mock_get_db.return_value = mock_db

        result = generate_validation_letters(
            client_id=999,
            collections=[{"creditor_name": "Test"}]
        )

        assert result["success"] is False
        assert "not found" in result["error"]

    @patch('services.debt_validation_service.get_db')
    def test_generate_letters_no_collections(self, mock_get_db):
        """Test error when no collections provided."""
        mock_db = MagicMock()
        mock_client = MagicMock()
        mock_client.id = 1
        mock_client.name = "Test Client"
        mock_db.query().filter().first.return_value = mock_client
        mock_get_db.return_value = mock_db

        result = generate_validation_letters(
            client_id=1,
            collections=[]
        )

        assert result["success"] is False
        assert "No collection accounts" in result["error"]

    @patch('services.debt_validation_service.get_db')
    @patch('services.debt_validation_service.os.makedirs')
    @patch.object(DebtValidationPDF, 'output')
    @patch('services.debt_validation_service.Document')
    def test_generate_letters_success(self, mock_document, mock_pdf_output,
                                      mock_makedirs, mock_get_db):
        """Test successful letter generation."""
        mock_db = MagicMock()
        mock_client = MagicMock()
        mock_client.id = 1
        mock_client.name = "Test Client"
        mock_client.first_name = "Test"
        mock_client.last_name = "Client"
        mock_client.address_street = "123 Test St"
        mock_client.address_city = "Test City"
        mock_client.address_state = "CA"
        mock_client.address_zip = "90001"
        mock_client.ssn_last_four = "1234"
        mock_client.date_of_birth = date(1980, 1, 1)
        mock_client.email = "test@test.com"
        mock_client.phone = "555-1234"

        mock_db.query().filter().first.return_value = mock_client
        mock_get_db.return_value = mock_db

        mock_doc = MagicMock()
        mock_document.return_value = mock_doc

        collections = [
            {
                "creditor_name": "ABC Collections",
                "creditor_address": "123 Collection St",
                "creditor_city_state_zip": "City, ST 12345",
                "original_creditor": "Original Bank",
                "account_number": "1234567890",
                "balance": 5000,
            }
        ]

        result = generate_validation_letters(
            client_id=1,
            collections=collections
        )

        assert result["success"] is True
        assert result["letters_generated"] == 1
        assert "ABC Collections" in result["collections"]

    @patch('services.debt_validation_service.get_db')
    @patch('services.debt_validation_service.os.makedirs')
    @patch.object(DebtValidationPDF, 'output')
    @patch('services.debt_validation_service.Document')
    def test_generate_letters_multiple_collections(self, mock_document, mock_pdf_output,
                                                   mock_makedirs, mock_get_db):
        """Test generating letters for multiple collections."""
        mock_db = MagicMock()
        mock_client = MagicMock()
        mock_client.id = 1
        mock_client.name = "Multi Collection Client"
        mock_client.first_name = "Multi"
        mock_client.last_name = "Client"
        mock_client.address_street = "456 Multi St"
        mock_client.address_city = "Multi City"
        mock_client.address_state = "TX"
        mock_client.address_zip = "75001"
        mock_client.ssn_last_four = "5678"
        mock_client.date_of_birth = None
        mock_client.email = None
        mock_client.phone = None

        mock_db.query().filter().first.return_value = mock_client
        mock_get_db.return_value = mock_db

        mock_doc = MagicMock()
        mock_document.return_value = mock_doc

        collections = [
            {"creditor_name": "Collector A", "account_number": "111", "balance": 1000},
            {"creditor_name": "Collector B", "account_number": "222", "balance": 2000},
            {"creditor_name": "Collector C", "account_number": "333", "balance": 3000},
        ]

        result = generate_validation_letters(
            client_id=1,
            collections=collections
        )

        assert result["success"] is True
        assert result["letters_generated"] == 3


# =============================================================================
# Tests for generate_validation_letter_single()
# =============================================================================

class TestGenerateValidationLetterSingle:
    """Test single letter generation wrapper."""

    @patch('services.debt_validation_service.generate_validation_letters')
    def test_single_letter_calls_main_function(self, mock_generate):
        """Test that single letter function calls main with collection list."""
        mock_generate.return_value = {"success": True, "letters_generated": 1}

        collection_info = {
            "creditor_name": "Single Collector",
            "account_number": "12345",
            "balance": 500,
        }

        result = generate_validation_letter_single(
            client_id=1,
            collection_info=collection_info
        )

        mock_generate.assert_called_once_with(1, collections=[collection_info])
        assert result["success"] is True


# =============================================================================
# Tests for auto_generate_validation_letters_from_analysis()
# =============================================================================

class TestAutoGenerateValidationLetters:
    """Test automatic validation letter generation from analysis."""

    @patch('services.debt_validation_service.get_db')
    def test_auto_generate_case_not_found(self, mock_get_db):
        """Test error when case is not found."""
        mock_db = MagicMock()
        mock_db.query().filter().first.return_value = None
        mock_get_db.return_value = mock_db

        result = auto_generate_validation_letters_from_analysis(case_id=999)

        assert result["success"] is False
        assert "Case not found" in result["error"]

    @patch('services.debt_validation_service.get_db')
    def test_auto_generate_client_not_found(self, mock_get_db):
        """Test error when client is not found."""
        mock_db = MagicMock()
        mock_case = MagicMock()
        mock_case.client_id = 999

        # First call returns case, second returns None for client
        mock_db.query().filter().first.side_effect = [mock_case, None]
        mock_get_db.return_value = mock_db

        result = auto_generate_validation_letters_from_analysis(case_id=1)

        assert result["success"] is False
        assert "Client not found" in result["error"]

    @patch('services.debt_validation_service.Violation')
    @patch('services.debt_validation_service.get_db')
    def test_auto_generate_no_collection_violations(self, mock_get_db, mock_violation_class):
        """Test when no collection-related violations are found."""
        mock_db = MagicMock()
        mock_case = MagicMock()
        mock_case.id = 1
        mock_case.client_id = 1

        mock_client = MagicMock()
        mock_client.id = 1
        mock_client.name = "Test Client"

        # Non-collection violation
        mock_violation = MagicMock()
        mock_violation.creditor_name = "Regular Bank"
        mock_violation.description = "Late payment reported"
        mock_violation.violation_type = "LATE_PAYMENT"
        mock_violation.account_number = "12345"

        mock_db.query().filter().first.side_effect = [mock_case, mock_client]
        mock_db.query().filter().all.return_value = [mock_violation]
        mock_get_db.return_value = mock_db

        result = auto_generate_validation_letters_from_analysis(case_id=1)

        assert result["success"] is True
        assert result["letters_generated"] == 0
        assert "No collection accounts" in result["message"]

    @patch('services.debt_validation_service.Violation')
    @patch('services.debt_validation_service.get_db')
    @patch('services.debt_validation_service.generate_validation_letters')
    def test_auto_generate_with_collections(self, mock_gen_letters, mock_get_db, mock_violation_class):
        """Test auto-generation with collection violations found."""
        mock_db = MagicMock()
        mock_case = MagicMock()
        mock_case.id = 1
        mock_case.client_id = 1

        mock_client = MagicMock()
        mock_client.id = 1

        # Collection violation
        mock_violation = MagicMock()
        mock_violation.creditor_name = "ABC Collections"
        mock_violation.description = "Debt collection account"
        mock_violation.violation_type = "COLLECTION_ACCOUNT"
        mock_violation.account_number = "COL12345"

        mock_db.query().filter().first.side_effect = [mock_case, mock_client]
        mock_db.query().filter().all.return_value = [mock_violation]
        mock_get_db.return_value = mock_db

        mock_gen_letters.return_value = {"success": True, "letters_generated": 1}

        result = auto_generate_validation_letters_from_analysis(case_id=1)

        mock_gen_letters.assert_called_once()
        assert result["success"] is True

    @patch('services.debt_validation_service.Violation')
    @patch('services.debt_validation_service.get_db')
    @patch('services.debt_validation_service.generate_validation_letters')
    def test_auto_generate_deduplicates_creditors(self, mock_gen_letters, mock_get_db, mock_violation_class):
        """Test that duplicate creditors are not included twice."""
        mock_db = MagicMock()
        mock_case = MagicMock()
        mock_case.id = 1
        mock_case.client_id = 1

        mock_client = MagicMock()
        mock_client.id = 1

        # Multiple violations for same creditor
        mock_v1 = MagicMock()
        mock_v1.creditor_name = "Same Collections"
        mock_v1.description = "Collection account"
        mock_v1.violation_type = "COLLECTION"
        mock_v1.account_number = "111"

        mock_v2 = MagicMock()
        mock_v2.creditor_name = "Same Collections"  # Same creditor
        mock_v2.description = "Another collection issue"
        mock_v2.violation_type = "DEBT_COLLECTOR"
        mock_v2.account_number = "222"

        mock_db.query().filter().first.side_effect = [mock_case, mock_client]
        mock_db.query().filter().all.return_value = [mock_v1, mock_v2]
        mock_get_db.return_value = mock_db

        mock_gen_letters.return_value = {"success": True, "letters_generated": 1}

        result = auto_generate_validation_letters_from_analysis(case_id=1)

        # Should only generate 1 letter (deduplicated)
        call_args = mock_gen_letters.call_args
        collections = call_args[1]["collections"]
        assert len(collections) == 1


# =============================================================================
# Tests for get_common_collection_agencies()
# =============================================================================

class TestGetCommonCollectionAgencies:
    """Test common collection agencies list."""

    def test_returns_list(self):
        """Test that function returns a list."""
        agencies = get_common_collection_agencies()
        assert isinstance(agencies, list)

    def test_agencies_have_required_fields(self):
        """Test that each agency has required fields."""
        agencies = get_common_collection_agencies()

        for agency in agencies:
            assert "name" in agency
            assert "address" in agency
            assert "city_state_zip" in agency

    def test_known_agencies_included(self):
        """Test that well-known collection agencies are included."""
        agencies = get_common_collection_agencies()
        agency_names = [a["name"] for a in agencies]

        assert "Portfolio Recovery Associates" in agency_names
        assert "Midland Credit Management" in agency_names
        assert "LVNV Funding LLC" in agency_names

    def test_agencies_count(self):
        """Test that a reasonable number of agencies are returned."""
        agencies = get_common_collection_agencies()
        assert len(agencies) >= 5  # At least 5 common agencies


# =============================================================================
# Tests for Edge Cases and Error Handling
# =============================================================================

class TestEdgeCasesAndErrorHandling:
    """Test edge cases and error handling."""

    def test_pii_with_unicode_characters(self):
        """Test PII extraction with unicode characters in name."""
        client = {
            "name": "Jose Garcia-Martinez",
            "address_street": "123 Calle Principal",
        }
        pii = get_client_pii(client)
        assert pii["name"] == "Jose Garcia-Martinez"

    def test_letter_with_zero_balance(self):
        """Test letter generation with zero balance."""
        client = {"name": "Zero Balance User"}
        collection_info = {
            "account_number": "12345",
            "balance": 0,
        }

        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["balance"] == "$0.00"

    def test_letter_with_negative_balance(self):
        """Test letter generation with negative balance (credit)."""
        client = {"name": "Credit Balance User"}
        collection_info = {
            "account_number": "12345",
            "balance": -100.50,
        }

        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["balance"] == "$-100.50"

    def test_letter_with_very_large_balance(self):
        """Test letter generation with very large balance."""
        client = {"name": "Big Debt User"}
        collection_info = {
            "account_number": "12345",
            "balance": 1234567890.99,
        }

        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["balance"] == "$1,234,567,890.99"

    def test_letter_with_special_characters_in_creditor(self):
        """Test letter with special characters in creditor name."""
        client = {"name": "Test User"}
        collection_info = {
            "creditor_name": "ABC & XYZ Collections, LLC.",
            "account_number": "12345",
            "balance": 100,
        }

        letter = generate_debt_validation_letter(client, collection_info)
        assert letter["creditor_name"] == "ABC & XYZ Collections, LLC."

    def test_collection_agency_template_structure(self):
        """Test the collection agency template constant."""
        assert "generic" in COLLECTION_AGENCY_TEMPLATE
        template = COLLECTION_AGENCY_TEMPLATE["generic"]
        assert "name" in template
        assert "address" in template
        assert "city_state_zip" in template

    @patch('services.debt_validation_service.get_db')
    def test_generate_letters_exception_handling(self, mock_get_db):
        """Test exception handling in generate_validation_letters."""
        mock_db = MagicMock()
        mock_db.query().filter().first.side_effect = Exception("Database error")
        mock_get_db.return_value = mock_db

        result = generate_validation_letters(
            client_id=1,
            collections=[{"creditor_name": "Test"}]
        )

        assert result["success"] is False
        assert "Database error" in result["error"]

    @patch('services.debt_validation_service.get_db')
    def test_auto_generate_exception_handling(self, mock_get_db):
        """Test exception handling in auto_generate_validation_letters_from_analysis."""
        mock_db = MagicMock()
        mock_db.query().filter().first.side_effect = Exception("Query failed")
        mock_get_db.return_value = mock_db

        result = auto_generate_validation_letters_from_analysis(case_id=1)

        assert result["success"] is False
        assert "Query failed" in result["error"]


# =============================================================================
# Tests for _add_validation_letter_to_docx()
# =============================================================================

class TestAddValidationLetterToDocx:
    """Test adding validation letter to Word document."""

    @patch('services.debt_validation_service.Document')
    def test_add_letter_to_docx_basic(self, mock_document_class):
        """Test adding a basic letter to Word document."""
        from docx import Document

        mock_doc = MagicMock()
        mock_paragraph = MagicMock()
        mock_run = MagicMock()
        mock_font = MagicMock()
        mock_run.font = mock_font
        mock_paragraph.add_run.return_value = mock_run
        mock_paragraph.paragraph_format = MagicMock()
        mock_doc.add_paragraph.return_value = mock_paragraph

        letter_content = {
            "date": "January 15, 2024",
            "client_name": "Test Client",
            "client_address": "123 Test St",
            "client_city_state_zip": "Test City, CA 90001",
            "creditor_name": "Test Collector",
            "creditor_address": "456 Collect Ave",
            "creditor_city_state_zip": "Collect City, TX 75001",
            "original_creditor": "Original Bank",
            "account_number": "XXXX1234",
            "balance": "$1,000.00",
            "ssn_last4": "1234",
            "reference_number": "",
        }

        _add_validation_letter_to_docx(mock_doc, letter_content, add_page_break=True)

        # Verify add_paragraph was called multiple times
        assert mock_doc.add_paragraph.called
        # Verify page break was added
        assert mock_doc.add_page_break.called

    @patch('services.debt_validation_service.Document')
    def test_add_letter_to_docx_no_page_break(self, mock_document_class):
        """Test adding letter without page break."""
        mock_doc = MagicMock()
        mock_paragraph = MagicMock()
        mock_run = MagicMock()
        mock_font = MagicMock()
        mock_run.font = mock_font
        mock_paragraph.add_run.return_value = mock_run
        mock_paragraph.paragraph_format = MagicMock()
        mock_doc.add_paragraph.return_value = mock_paragraph

        letter_content = {
            "date": "January 15, 2024",
            "client_name": "Test Client",
            "client_address": "123 Test St",
            "client_city_state_zip": "Test City, CA 90001",
            "creditor_name": "Test Collector",
            "creditor_address": "456 Collect Ave",
            "creditor_city_state_zip": "Collect City, TX 75001",
            "original_creditor": "Original Bank",
            "account_number": "XXXX1234",
            "balance": "$1,000.00",
            "ssn_last4": "1234",
            "reference_number": "",
        }

        _add_validation_letter_to_docx(mock_doc, letter_content, add_page_break=False)

        # Verify page break was NOT added
        assert not mock_doc.add_page_break.called
