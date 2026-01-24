"""
PII Correction Letter Service
Generates FCRA-compliant PII correction letters for the Big 3 CRAs.
Used when client has incorrect personal identifying information on their credit reports.
Uses fpdf2 for PDF generation and python-docx for Word document generation.
"""

import os
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt
from fpdf import FPDF

from database import Client, get_db

# Big 3 CRA addresses for PII correction letters
CRA_ADDRESSES = {
    "Equifax": {
        "name": "Equifax Information Services LLC",
        "dispute_address": "P.O. Box 740256",
        "city": "Atlanta",
        "state": "GA",
        "zip": "30374",
        "phone": "1-800-685-1111",
    },
    "Experian": {
        "name": "Experian",
        "dispute_address": "P.O. Box 4500",
        "city": "Allen",
        "state": "TX",
        "zip": "75013",
        "phone": "1-888-397-3742",
    },
    "TransUnion": {
        "name": "TransUnion LLC",
        "dispute_address": "P.O. Box 2000",
        "city": "Chester",
        "state": "PA",
        "zip": "19016",
        "phone": "1-800-916-8800",
    },
}

# PII types that can be corrected
PII_TYPES = [
    "name",
    "address",
    "ssn",
    "dob",
    "phone",
    "employer",
]


class PIICorrectionPDF(FPDF):
    """Custom PDF class for PII correction letters"""

    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=25)

    def header(self):
        pass

    def footer(self):
        self.set_y(-20)
        self.set_font("Arial", "I", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")


def get_client_pii(client) -> Dict[str, Any]:
    """Extract PII from client object or dict"""
    if isinstance(client, dict):
        name = (
            client.get("name", "")
            or f"{client.get('first_name', '')} {client.get('last_name', '')}".strip()
        )
        return {
            "name": name,
            "first_name": client.get("first_name", ""),
            "last_name": client.get("last_name", ""),
            "address_street": client.get("address_street", ""),
            "address_city": client.get("address_city", ""),
            "address_state": client.get("address_state", ""),
            "address_zip": client.get("address_zip", ""),
            "ssn_last_four": client.get("ssn_last_four", "XXXX"),
            "date_of_birth": client.get("date_of_birth", None),
            "email": client.get("email", ""),
            "phone": client.get("phone", ""),
            "employer": client.get("employer", "")
            or client.get("employer_company", ""),
        }
    else:
        name = client.name or ""
        if not name and client.first_name and client.last_name:
            name = f"{client.first_name} {client.last_name}"
        return {
            "name": name,
            "first_name": client.first_name or "",
            "last_name": client.last_name or "",
            "address_street": client.address_street or "",
            "address_city": client.address_city or "",
            "address_state": client.address_state or "",
            "address_zip": client.address_zip or "",
            "ssn_last_four": client.ssn_last_four or "XXXX",
            "date_of_birth": client.date_of_birth,
            "email": client.email or "",
            "phone": client.phone or "",
            "employer": getattr(client, "employer", "")
            or getattr(client, "employer_company", "")
            or "",
        }


def generate_pii_correction_letter(
    client,
    bureau_name: str,
    incorrect_pii: Dict[str, List[str]],
    correct_pii: Optional[Dict[str, str]] = None,
    case_number: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Generate a single PII correction letter for a specific bureau.

    Args:
        client: Client object or dict with client information
        bureau_name: Name of the bureau (Equifax, Experian, TransUnion)
        incorrect_pii: Dict of PII type to list of incorrect values to DELETE
            Example: {
                "names": ["JOHN DOE JR", "J DOE"],
                "addresses": ["123 OLD ST, OLDTOWN, CA 90000"],
                "phones": ["555-123-4567"],
                "employers": ["OLD COMPANY INC"]
            }
        correct_pii: Optional dict of correct values (uses client data if not provided)
            Example: {
                "name": "John Doe",
                "address": "456 New St, Newtown, CA 90001",
                "ssn": "XXX-XX-1234",
                "dob": "01/15/1985",
                "phone": "555-987-6543",
                "employer": "Current Company LLC"
            }
        case_number: Optional case reference number

    Returns:
        dict: Letter content for PDF generation
    """
    if bureau_name not in CRA_ADDRESSES:
        raise ValueError(
            f"Unknown bureau: {bureau_name}. Must be one of: {list(CRA_ADDRESSES.keys())}"
        )

    bureau_info = CRA_ADDRESSES[bureau_name]
    client_pii = get_client_pii(client)
    today = datetime.now().strftime("%B %d, %Y")

    # Build correct PII from client data if not provided
    if correct_pii is None:
        correct_pii = {}

    # Use client data as defaults for correct PII
    if "name" not in correct_pii:
        correct_pii["name"] = client_pii["name"]
    if "address" not in correct_pii:
        addr_parts = [
            client_pii["address_street"],
            f"{client_pii['address_city']}, {client_pii['address_state']} {client_pii['address_zip']}",
        ]
        correct_pii["address"] = ", ".join(p for p in addr_parts if p.strip())
    if "ssn" not in correct_pii:
        correct_pii["ssn"] = f"XXX-XX-{client_pii['ssn_last_four']}"
    if "dob" not in correct_pii:
        if client_pii["date_of_birth"]:
            if hasattr(client_pii["date_of_birth"], "strftime"):
                correct_pii["dob"] = client_pii["date_of_birth"].strftime("%m/%d/%Y")
            else:
                correct_pii["dob"] = str(client_pii["date_of_birth"])
        else:
            correct_pii["dob"] = "[DATE OF BIRTH]"
    if "phone" not in correct_pii:
        correct_pii["phone"] = client_pii["phone"] or "REMOVE ALL"
    if "employer" not in correct_pii:
        correct_pii["employer"] = client_pii["employer"] or "REMOVE ALL"

    # Generate case number if not provided
    if not case_number:
        case_number = (
            f"PII-{datetime.now().strftime('%Y%m%d')}-{client_pii['ssn_last_four']}"
        )

    bureau_full_address = f"{bureau_info['dispute_address']}\n{bureau_info['city']}, {bureau_info['state']} {bureau_info['zip']}"

    letter_content = {
        "date": today,
        "bureau_name": bureau_info["name"],
        "bureau_address": bureau_full_address,
        "client_name": client_pii["name"],
        "client_address": correct_pii["address"],
        "ssn": correct_pii["ssn"],
        "dob": correct_pii["dob"],
        "case_number": case_number,
        "incorrect_pii": incorrect_pii,
        "correct_pii": correct_pii,
    }

    return letter_content


def _add_pii_letter_to_pdf(pdf: FPDF, letter_content: Dict[str, Any]) -> None:
    """Add a single PII correction letter to the PDF document"""
    pdf.add_page()

    # Consumer address block
    pdf.set_font("Arial", "", 11)
    pdf.cell(0, 6, letter_content["client_name"], ln=True)
    for line in letter_content["client_address"].split(", "):
        if line.strip():
            pdf.cell(0, 6, line.strip(), ln=True)
    pdf.ln(6)

    # Date
    pdf.cell(0, 6, letter_content["date"], ln=True)
    pdf.ln(6)

    # Bureau address block
    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 6, letter_content["bureau_name"], ln=True)
    pdf.set_font("Arial", "", 11)
    for line in letter_content["bureau_address"].split("\n"):
        if line.strip():
            pdf.cell(0, 6, line.strip(), ln=True)
    pdf.ln(4)
    pdf.cell(0, 6, "SENT VIA CERTIFIED MAIL", ln=True)
    pdf.ln(8)

    # Subject line
    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 6, "=" * 70, ln=True)
    pdf.cell(
        0, 8, "RE: FORMAL DEMAND TO CORRECT PERSONAL IDENTIFYING INFORMATION", ln=True
    )
    pdf.cell(
        0,
        6,
        "FCRA Sections 1681e(b), 1681i - INACCURATE PII VIOLATES FEDERAL LAW",
        ln=True,
    )
    pdf.cell(0, 6, "=" * 70, ln=True)
    pdf.ln(4)

    # Consumer info block
    pdf.set_font("Arial", "", 11)
    pdf.cell(40, 6, "Consumer:", 0)
    pdf.cell(0, 6, letter_content["client_name"], ln=True)
    pdf.cell(40, 6, "SSN:", 0)
    pdf.cell(0, 6, letter_content["ssn"], ln=True)
    pdf.cell(40, 6, "Case Reference:", 0)
    pdf.cell(0, 6, letter_content["case_number"], ln=True)
    pdf.ln(6)

    # Greeting
    pdf.cell(0, 6, "Dear Sir or Madam:", ln=True)
    pdf.ln(4)

    # Intro paragraph
    intro = (
        f"I have reviewed my credit file from {letter_content['bureau_name']} and discovered "
        "INACCURATE PERSONAL IDENTIFYING INFORMATION in violation of FCRA Section 1681e(b)."
    )
    pdf.multi_cell(0, 6, intro)
    pdf.ln(4)
    pdf.set_font("Arial", "B", 11)
    pdf.cell(
        0, 6, "This is not a request. This is a DEMAND backed by federal law.", ln=True
    )
    pdf.set_font("Arial", "", 11)
    pdf.ln(6)

    # Section 1: Incorrect PII to DELETE
    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 6, "-" * 70, ln=True)
    pdf.cell(0, 8, "SECTION 1: INACCURATE PII IN YOUR FILE - DELETE THESE", ln=True)
    pdf.cell(0, 6, "-" * 70, ln=True)
    pdf.set_font("Arial", "", 11)
    pdf.ln(4)

    incorrect = letter_content["incorrect_pii"]

    if incorrect.get("names"):
        pdf.set_font("Arial", "B", 11)
        pdf.cell(0, 6, "INCORRECT NAMES:", ln=True)
        pdf.set_font("Arial", "", 11)
        for name in incorrect["names"]:
            pdf.cell(0, 6, f"  X {name} - DELETE", ln=True)
        pdf.ln(2)

    if incorrect.get("addresses"):
        pdf.set_font("Arial", "B", 11)
        pdf.cell(0, 6, "INCORRECT ADDRESSES:", ln=True)
        pdf.set_font("Arial", "", 11)
        for addr in incorrect["addresses"]:
            pdf.cell(0, 6, f"  X {addr} - DELETE", ln=True)
        pdf.ln(2)

    if incorrect.get("phones"):
        pdf.set_font("Arial", "B", 11)
        pdf.cell(0, 6, "INCORRECT PHONE NUMBERS:", ln=True)
        pdf.set_font("Arial", "", 11)
        for phone in incorrect["phones"]:
            pdf.cell(0, 6, f"  X {phone} - DELETE", ln=True)
        pdf.ln(2)

    if incorrect.get("employers"):
        pdf.set_font("Arial", "B", 11)
        pdf.cell(0, 6, "INCORRECT EMPLOYERS:", ln=True)
        pdf.set_font("Arial", "", 11)
        for emp in incorrect["employers"]:
            pdf.cell(0, 6, f"  X {emp} - DELETE", ln=True)
        pdf.ln(2)

    if incorrect.get("ssn_variations"):
        pdf.set_font("Arial", "B", 11)
        pdf.cell(0, 6, "INCORRECT SSN VARIATIONS:", ln=True)
        pdf.set_font("Arial", "", 11)
        for ssn in incorrect["ssn_variations"]:
            pdf.cell(0, 6, f"  X {ssn} - DELETE", ln=True)
        pdf.ln(2)

    pdf.ln(4)

    # Section 2: Correct PII to KEEP
    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 6, "-" * 70, ln=True)
    pdf.cell(0, 8, "SECTION 2: CORRECT INFORMATION - KEEP ONLY THESE", ln=True)
    pdf.cell(0, 6, "-" * 70, ln=True)
    pdf.set_font("Arial", "", 11)
    pdf.ln(4)

    correct = letter_content["correct_pii"]
    pdf.cell(0, 6, f"NAME: {correct['name']} - DELETE ALL OTHERS", ln=True)
    pdf.cell(0, 6, f"ADDRESS: {correct['address']} - DELETE ALL OTHERS", ln=True)
    pdf.cell(0, 6, f"SSN: {correct['ssn']} - DELETE ALL VARIATIONS", ln=True)
    pdf.cell(0, 6, f"DOB: {correct['dob']}", ln=True)
    pdf.cell(0, 6, f"PHONE: {correct['phone']} - DELETE ALL OTHERS", ln=True)
    pdf.cell(0, 6, f"EMPLOYER: {correct['employer']} - DELETE ALL OTHERS", ln=True)
    pdf.ln(6)

    # Section 3: Why This Matters
    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 6, "-" * 70, ln=True)
    pdf.cell(0, 8, "SECTION 3: WHY THIS MATTERS - LEGAL VIOLATIONS", ln=True)
    pdf.cell(0, 6, "-" * 70, ln=True)
    pdf.set_font("Arial", "", 11)
    pdf.ln(4)

    why_matters = (
        "Inaccurate PII violates FCRA Section 1681e(b) and causes direct harm:\n\n"
        "1. Mixed file risk - Incorrect PII can mix my file with another consumer\n"
        "2. Identity theft vulnerability\n"
        "3. Inaccurate credit reporting\n"
        "4. Employment and housing impact\n\n"
        "EVERY DAY you maintain inaccurate PII = continuing Section 1681e(b) violation."
    )
    pdf.multi_cell(0, 6, why_matters)
    pdf.ln(6)

    # Section 4: Damages Exposure
    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 6, "-" * 70, ln=True)
    pdf.cell(0, 8, "SECTION 4: DAMAGES EXPOSURE", ln=True)
    pdf.cell(0, 6, "-" * 70, ln=True)
    pdf.set_font("Arial", "", 11)
    pdf.ln(4)

    damages_text = (
        "Failure to maintain accuracy: $100-$1,000 per violation\n"
        "Failure to correct after notice: $100-$1,000 per violation\n"
        "Mixed file (if applicable): $1,000+\n"
        "Punitive (if willful): 2-4x statutory\n"
        "Attorney fees: $15,000+\n\n"
        "You are on notice. Continued inaccuracy after this letter = willfulness under Safeco v. Burr."
    )
    pdf.multi_cell(0, 6, damages_text)
    pdf.ln(6)

    # Section 5: Demand
    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 6, "-" * 70, ln=True)
    pdf.cell(0, 8, "SECTION 5: DEMAND", ln=True)
    pdf.cell(0, 6, "-" * 70, ln=True)
    pdf.set_font("Arial", "", 11)
    pdf.ln(4)

    demand = (
        f"I DEMAND {letter_content['bureau_name']}:\n\n"
        "1. DELETE all incorrect PII listed in Section 1\n"
        "2. CONFIRM corrections in writing within 30 days\n"
        "3. PROVIDE updated disclosure showing ONLY correct PII\n\n"
        "Non-compliance = CFPB, FTC, State AG complaints + FCRA litigation."
    )
    pdf.multi_cell(0, 6, demand)
    pdf.ln(8)

    # Signature block
    pdf.cell(0, 6, "Sincerely,", ln=True)
    pdf.ln(16)
    pdf.cell(0, 6, "_" * 40, ln=True)
    pdf.cell(0, 6, letter_content["client_name"], ln=True)
    pdf.ln(4)
    pdf.cell(0, 6, f"Date: {letter_content['date']}", ln=True)
    pdf.ln(6)

    # Enclosures
    pdf.set_font("Arial", "B", 11)
    pdf.cell(0, 6, "ENCLOSURES:", ln=True)
    pdf.set_font("Arial", "", 11)
    pdf.cell(0, 6, "- Copy of government-issued photo ID", ln=True)
    pdf.cell(0, 6, "- Copy of Social Security Card", ln=True)
    pdf.cell(
        0, 6, "- Proof of current address (utility bill or bank statement)", ln=True
    )


def _add_pii_letter_to_docx(
    doc: Document, letter_content: Dict[str, Any], add_page_break: bool = True
) -> None:
    """Add a single PII correction letter to the Word document"""

    def add_para(text: str, bold: bool = False, size: int = 11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        return p

    # Consumer address
    add_para(letter_content["client_name"])
    for line in letter_content["client_address"].split(", "):
        if line.strip():
            add_para(line.strip())
    doc.add_paragraph()

    # Date
    add_para(letter_content["date"])
    doc.add_paragraph()

    # Bureau address
    add_para(letter_content["bureau_name"], bold=True)
    for line in letter_content["bureau_address"].split("\n"):
        if line.strip():
            add_para(line.strip())
    add_para("SENT VIA CERTIFIED MAIL", bold=True)
    doc.add_paragraph()

    # Subject
    add_para("=" * 70)
    add_para("RE: FORMAL DEMAND TO CORRECT PERSONAL IDENTIFYING INFORMATION", bold=True)
    add_para(
        "FCRA Sections 1681e(b), 1681i - INACCURATE PII VIOLATES FEDERAL LAW", bold=True
    )
    add_para("=" * 70)
    doc.add_paragraph()

    # Consumer info
    add_para(f"Consumer: {letter_content['client_name']}")
    add_para(f"SSN: {letter_content['ssn']}")
    add_para(f"Case Reference: {letter_content['case_number']}")
    doc.add_paragraph()

    # Body
    add_para("Dear Sir or Madam:")
    doc.add_paragraph()

    intro = (
        f"I have reviewed my credit file from {letter_content['bureau_name']} and discovered "
        "INACCURATE PERSONAL IDENTIFYING INFORMATION in violation of FCRA Section 1681e(b)."
    )
    add_para(intro)
    add_para(
        "This is not a request. This is a DEMAND backed by federal law.", bold=True
    )
    doc.add_paragraph()

    # Section 1
    add_para("-" * 70)
    add_para("SECTION 1: INACCURATE PII IN YOUR FILE - DELETE THESE", bold=True)
    add_para("-" * 70)

    incorrect = letter_content["incorrect_pii"]
    if incorrect.get("names"):
        add_para("INCORRECT NAMES:", bold=True)
        for name in incorrect["names"]:
            add_para(f"  X {name} - DELETE")
    if incorrect.get("addresses"):
        add_para("INCORRECT ADDRESSES:", bold=True)
        for addr in incorrect["addresses"]:
            add_para(f"  X {addr} - DELETE")
    if incorrect.get("phones"):
        add_para("INCORRECT PHONE NUMBERS:", bold=True)
        for phone in incorrect["phones"]:
            add_para(f"  X {phone} - DELETE")
    if incorrect.get("employers"):
        add_para("INCORRECT EMPLOYERS:", bold=True)
        for emp in incorrect["employers"]:
            add_para(f"  X {emp} - DELETE")
    doc.add_paragraph()

    # Section 2
    add_para("-" * 70)
    add_para("SECTION 2: CORRECT INFORMATION - KEEP ONLY THESE", bold=True)
    add_para("-" * 70)

    correct = letter_content["correct_pii"]
    add_para(f"NAME: {correct['name']} - DELETE ALL OTHERS")
    add_para(f"ADDRESS: {correct['address']} - DELETE ALL OTHERS")
    add_para(f"SSN: {correct['ssn']} - DELETE ALL VARIATIONS")
    add_para(f"DOB: {correct['dob']}")
    add_para(f"PHONE: {correct['phone']} - DELETE ALL OTHERS")
    add_para(f"EMPLOYER: {correct['employer']} - DELETE ALL OTHERS")
    doc.add_paragraph()

    # Section 3
    add_para("-" * 70)
    add_para("SECTION 3: WHY THIS MATTERS - LEGAL VIOLATIONS", bold=True)
    add_para("-" * 70)
    add_para("Inaccurate PII violates FCRA Section 1681e(b) and causes direct harm:")
    add_para("1. Mixed file risk - Incorrect PII can mix my file with another consumer")
    add_para("2. Identity theft vulnerability")
    add_para("3. Inaccurate credit reporting")
    add_para("4. Employment and housing impact")
    add_para(
        "EVERY DAY you maintain inaccurate PII = continuing Section 1681e(b) violation.",
        bold=True,
    )
    doc.add_paragraph()

    # Section 4
    add_para("-" * 70)
    add_para("SECTION 4: DAMAGES EXPOSURE", bold=True)
    add_para("-" * 70)
    add_para("Failure to maintain accuracy: $100-$1,000 per violation")
    add_para("Failure to correct after notice: $100-$1,000 per violation")
    add_para("Mixed file (if applicable): $1,000+")
    add_para("Punitive (if willful): 2-4x statutory")
    add_para("Attorney fees: $15,000+")
    add_para(
        "You are on notice. Continued inaccuracy after this letter = willfulness under Safeco v. Burr.",
        bold=True,
    )
    doc.add_paragraph()

    # Section 5
    add_para("-" * 70)
    add_para("SECTION 5: DEMAND", bold=True)
    add_para("-" * 70)
    add_para(f"I DEMAND {letter_content['bureau_name']}:")
    add_para("1. DELETE all incorrect PII listed in Section 1")
    add_para("2. CONFIRM corrections in writing within 30 days")
    add_para("3. PROVIDE updated disclosure showing ONLY correct PII")
    add_para(
        "Non-compliance = CFPB, FTC, State AG complaints + FCRA litigation.", bold=True
    )
    doc.add_paragraph()

    # Signature
    add_para("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph()
    add_para("_" * 40)
    add_para(letter_content["client_name"])
    add_para(f"Date: {letter_content['date']}")
    doc.add_paragraph()

    # Enclosures
    add_para("ENCLOSURES:", bold=True)
    add_para("- Copy of government-issued photo ID")
    add_para("- Copy of Social Security Card")
    add_para("- Proof of current address (utility bill or bank statement)")

    if add_page_break:
        doc.add_page_break()


def generate_pii_correction_letters(
    client_id: int,
    incorrect_pii: Dict[str, List[str]],
    correct_pii: Optional[Dict[str, str]] = None,
    bureaus: Optional[List[str]] = None,
    case_number: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Generate PII correction letters for specified bureaus (or all 3 if not specified).

    Args:
        client_id: ID of the client in the database
        incorrect_pii: Dict of PII type to list of incorrect values to DELETE
        correct_pii: Optional dict of correct values (uses client data if not provided)
        bureaus: List of bureau names, or None for all 3
        case_number: Optional case reference number

    Returns:
        dict: {
            'success': bool,
            'batch_id': str (UUID),
            'pdf_path': str,
            'docx_path': str,
            'bureaus_included': list,
            'error': str (if failed)
        }
    """
    db = get_db()

    try:
        client = db.query(Client).filter(Client.id == client_id).first()

        if not client:
            return {"success": False, "error": f"Client with ID {client_id} not found"}

        # Default to all 3 bureaus
        if bureaus is None:
            target_bureaus = list(CRA_ADDRESSES.keys())
        else:
            target_bureaus = [b for b in bureaus if b in CRA_ADDRESSES]
            if not target_bureaus:
                return {"success": False, "error": "No valid bureaus specified"}

        # Create output directory
        output_dir = f"static/client_uploads/{client_id}/pii_letters"
        os.makedirs(output_dir, exist_ok=True)

        batch_uuid = str(uuid.uuid4())
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        client_name_safe = "".join(
            c
            for c in (client.name or client.last_name or "client")
            if c.isalnum() or c in " _-"
        ).replace(" ", "_")

        pdf_filename = f"{client_name_safe}_PII_Correction_{timestamp}.pdf"
        pdf_path = os.path.join(output_dir, pdf_filename)

        pdf = PIICorrectionPDF()
        doc = Document()

        letter_contents = []
        for bureau_name in target_bureaus:
            letter_content = generate_pii_correction_letter(
                client,
                bureau_name,
                incorrect_pii,
                correct_pii,
                case_number,
            )
            letter_contents.append(letter_content)
            _add_pii_letter_to_pdf(pdf, letter_content)

        pdf.output(pdf_path)

        # Generate Word document
        for i, letter_content in enumerate(letter_contents):
            is_last = i == len(letter_contents) - 1
            _add_pii_letter_to_docx(doc, letter_content, add_page_break=not is_last)

        docx_filename = f"{client_name_safe}_PII_Correction_{timestamp}.docx"
        docx_path = os.path.join(output_dir, docx_filename)
        doc.save(docx_path)

        return {
            "success": True,
            "batch_id": batch_uuid,
            "pdf_path": pdf_path,
            "docx_path": docx_path,
            "bureaus_included": target_bureaus,
            "total_letters": len(target_bureaus),
        }

    except Exception as e:
        db.rollback()
        return {"success": False, "error": str(e)}
    finally:
        db.close()


def generate_pii_correction_for_bureau(
    client_id: int,
    bureau_name: str,
    incorrect_pii: Dict[str, List[str]],
    correct_pii: Optional[Dict[str, str]] = None,
) -> Dict[str, Any]:
    """
    Generate a single PII correction letter for one bureau.

    Args:
        client_id: ID of the client
        bureau_name: Name of bureau (Equifax, Experian, TransUnion)
        incorrect_pii: Dict of incorrect PII to delete
        correct_pii: Optional correct values

    Returns:
        dict with success status and file paths
    """
    return generate_pii_correction_letters(
        client_id=client_id,
        incorrect_pii=incorrect_pii,
        correct_pii=correct_pii,
        bureaus=[bureau_name],
    )


def get_pii_discrepancies_from_report(
    parsed_report: Dict[str, Any],
) -> Dict[str, List[str]]:
    """
    Extract PII discrepancies from a parsed credit report.

    This function analyzes a parsed credit report and identifies
    variations in PII that may need correction.

    Args:
        parsed_report: Dict containing parsed credit report data
            Expected keys: personal_info, addresses, employers, etc.

    Returns:
        Dict of PII type to list of variations found
    """
    discrepancies = {
        "names": [],
        "addresses": [],
        "phones": [],
        "employers": [],
        "ssn_variations": [],
    }

    personal_info = parsed_report.get("personal_info", {})

    # Check for name variations
    names = personal_info.get("names", [])
    if isinstance(names, list) and len(names) > 1:
        # First name is assumed correct, rest are variations
        discrepancies["names"] = names[1:]
    elif personal_info.get("aka_names"):
        discrepancies["names"] = personal_info.get("aka_names", [])

    # Check for address variations
    addresses = personal_info.get("addresses", [])
    if isinstance(addresses, list) and len(addresses) > 1:
        # Current address is first, previous are potential issues
        discrepancies["addresses"] = [
            addr.get("full_address", str(addr)) if isinstance(addr, dict) else str(addr)
            for addr in addresses[1:]
        ]

    # Check for phone variations
    phones = personal_info.get("phones", [])
    if isinstance(phones, list):
        discrepancies["phones"] = phones

    # Check for employer variations
    employers = personal_info.get("employers", [])
    if isinstance(employers, list) and len(employers) > 1:
        discrepancies["employers"] = employers[1:]

    return discrepancies


# Convenience function for quick PII-only dispute
def generate_pii_only_dispute(
    client_id: int, incorrect_items: Dict[str, List[str]]
) -> Dict[str, Any]:
    """
    Quick function to generate PII correction letters for a client with clean credit
    who only needs PII corrections (no negative items to dispute).

    Args:
        client_id: Client database ID
        incorrect_items: Dict with keys like 'names', 'addresses', 'phones', 'employers'
                        and values as lists of incorrect items to remove

    Returns:
        Result dict with success status and file paths
    """
    return generate_pii_correction_letters(
        client_id=client_id,
        incorrect_pii=incorrect_items,
    )
