"""
Freeze Letter Service
Generates bulk credit freeze letters for all 12 credit bureaus.
Uses fpdf2 for PDF generation and python-docx for Word document generation.
"""
import os
import uuid
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from database import get_db, Client, FreezeLetterBatch


CRA_ADDRESSES = {
    'Equifax': {
        'name': 'Equifax Information Services LLC',
        'freeze_address': 'P.O. Box 105788',
        'city': 'Atlanta',
        'state': 'GA',
        'zip': '30348',
        'phone': '1-800-685-1111',
        'website': 'www.equifax.com/personal/credit-report-services/credit-freeze/',
        'type': 'primary'
    },
    'Experian': {
        'name': 'Experian Security Freeze',
        'freeze_address': 'P.O. Box 9554',
        'city': 'Allen',
        'state': 'TX',
        'zip': '75013',
        'phone': '1-888-397-3742',
        'website': 'www.experian.com/freeze/center.html',
        'type': 'primary'
    },
    'TransUnion': {
        'name': 'TransUnion LLC',
        'freeze_address': 'P.O. Box 160',
        'city': 'Woodlyn',
        'state': 'PA',
        'zip': '19094',
        'phone': '1-888-909-8872',
        'website': 'www.transunion.com/credit-freeze',
        'type': 'primary'
    },
    'Innovis': {
        'name': 'Innovis Consumer Assistance',
        'freeze_address': 'P.O. Box 26',
        'city': 'Pittsburgh',
        'state': 'PA',
        'zip': '15230-0026',
        'phone': '1-800-540-2505',
        'website': 'www.innovis.com/personal/securityFreeze',
        'type': 'secondary'
    },
    'ChexSystems': {
        'name': 'ChexSystems Inc.',
        'freeze_address': 'Attn: Security Freeze, P.O. Box 583399',
        'city': 'Minneapolis',
        'state': 'MN',
        'zip': '55458',
        'phone': '1-800-887-7652',
        'website': 'www.chexsystems.com',
        'type': 'secondary'
    },
    'Clarity Services Inc': {
        'name': 'Clarity Services Inc.',
        'freeze_address': 'P.O. Box 5717',
        'city': 'Clearwater',
        'state': 'FL',
        'zip': '33758',
        'phone': '1-866-390-3118',
        'website': 'www.clarityservices.com',
        'type': 'secondary'
    },
    'LexisNexis': {
        'name': 'LexisNexis Consumer Center',
        'freeze_address': 'P.O. Box 105108',
        'city': 'Atlanta',
        'state': 'GA',
        'zip': '30348-5108',
        'phone': '1-888-497-0011',
        'website': 'consumer.risk.lexisnexis.com',
        'type': 'secondary'
    },
    'CoreLogic Teletrack': {
        'name': 'CoreLogic Teletrack',
        'freeze_address': 'P.O. Box 509124',
        'city': 'San Diego',
        'state': 'CA',
        'zip': '92150',
        'phone': '1-877-309-5226',
        'website': 'www.corelogic.com',
        'type': 'secondary'
    },
    'Factor Trust Inc': {
        'name': 'FactorTrust Inc.',
        'freeze_address': '2930 Powers Ferry Road SE, Suite 301',
        'city': 'Atlanta',
        'state': 'GA',
        'zip': '30339',
        'phone': '1-844-300-1280',
        'website': 'www.factortrust.com',
        'type': 'secondary'
    },
    'MicroBilt/PRBC': {
        'name': 'MicroBilt Corporation / PRBC',
        'freeze_address': '1640 Airport Road, Suite 115',
        'city': 'Kennesaw',
        'state': 'GA',
        'zip': '30144',
        'phone': '1-866-205-7921',
        'website': 'www.microbilt.com',
        'type': 'secondary'
    },
    'LexisNexis Risk Solutions': {
        'name': 'LexisNexis Risk Solutions',
        'freeze_address': 'Consumer Center, P.O. Box 105108',
        'city': 'Atlanta',
        'state': 'GA',
        'zip': '30348-5108',
        'phone': '1-888-497-0011',
        'website': 'risk.lexisnexis.com',
        'type': 'secondary'
    },
    'DataX Ltd': {
        'name': 'DataX Ltd',
        'freeze_address': '110 E. Center Street, Suite 200',
        'city': 'Madison',
        'state': 'SD',
        'zip': '57042',
        'phone': '1-800-295-4790',
        'website': 'www.dataxltd.com',
        'type': 'secondary'
    },
}


def get_all_bureau_addresses():
    """
    Return dict of all 12 bureaus with their freeze addresses and phone numbers.
    
    Returns:
        dict: All bureau information including freeze addresses, phone numbers, and websites
    """
    return CRA_ADDRESSES.copy()


def get_primary_bureaus():
    """Return only the 3 primary CRAs (Equifax, Experian, TransUnion)"""
    return {k: v for k, v in CRA_ADDRESSES.items() if v.get('type') == 'primary'}


def get_secondary_bureaus():
    """Return only the 9 secondary bureaus"""
    return {k: v for k, v in CRA_ADDRESSES.items() if v.get('type') == 'secondary'}


class FreezeLetterPDF(FPDF):
    """Custom PDF class for freeze letters - clean format for bureau correspondence"""
    
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=25)
    
    def header(self):
        pass
    
    def footer(self):
        self.set_y(-20)
        self.set_font('Arial', 'I', 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f'Page {self.page_no()}', align='C')


def generate_single_freeze_letter(client, bureau_name, bureau_address):
    """
    Create formal letter requesting credit freeze for a single bureau.
    
    Args:
        client: Client object or dict with client information
        bureau_name: Name of the bureau (e.g., 'Equifax')
        bureau_address: Dict with bureau address information
    
    Returns:
        FPDF page content (adds to existing PDF)
    """
    if isinstance(client, dict):
        client_name = client.get('name', '')
        first_name = client.get('first_name', '')
        last_name = client.get('last_name', '')
        address_street = client.get('address_street', '')
        address_city = client.get('address_city', '')
        address_state = client.get('address_state', '')
        address_zip = client.get('address_zip', '')
        ssn_last_four = client.get('ssn_last_four', 'XXXX')
        date_of_birth = client.get('date_of_birth', None)
    else:
        client_name = client.name or ''
        first_name = client.first_name or ''
        last_name = client.last_name or ''
        address_street = client.address_street or ''
        address_city = client.address_city or ''
        address_state = client.address_state or ''
        address_zip = client.address_zip or ''
        ssn_last_four = client.ssn_last_four or 'XXXX'
        date_of_birth = client.date_of_birth
    
    if not client_name and first_name and last_name:
        client_name = f"{first_name} {last_name}"
    
    if date_of_birth:
        if hasattr(date_of_birth, 'strftime'):
            dob_formatted = date_of_birth.strftime('%B %d, %Y')
        else:
            dob_formatted = str(date_of_birth)
    else:
        dob_formatted = '[DATE OF BIRTH]'
    
    client_full_address = f"{address_street}, {address_city}, {address_state} {address_zip}".strip(', ')
    if not client_full_address or client_full_address == ', , ':
        client_full_address = '[CLIENT ADDRESS]'
    
    bureau_full_address = f"{bureau_address.get('freeze_address', '')}\n{bureau_address.get('city', '')}, {bureau_address.get('state', '')} {bureau_address.get('zip', '')}"
    
    today = datetime.now().strftime('%B %d, %Y')
    
    letter_content = {
        'date': today,
        'bureau_name': bureau_address.get('name', bureau_name),
        'bureau_address': bureau_full_address,
        'client_name': client_name,
        'client_address': client_full_address,
        'dob': dob_formatted,
        'ssn_last4': ssn_last_four,
        'bureau_phone': bureau_address.get('phone', ''),
    }
    
    return letter_content


def _add_letter_to_pdf(pdf, letter_content):
    """Add a single freeze letter to the PDF document"""
    pdf.add_page()
    
    pdf.set_font('Arial', '', 11)
    pdf.cell(0, 8, letter_content['date'], ln=True)
    pdf.ln(8)
    
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 6, letter_content['bureau_name'], ln=True)
    pdf.set_font('Arial', '', 11)
    for line in letter_content['bureau_address'].split('\n'):
        if line.strip():
            pdf.cell(0, 6, line.strip(), ln=True)
    
    pdf.ln(10)
    
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 8, 'RE: Request for Security Freeze on Credit File', ln=True)
    pdf.ln(6)
    
    pdf.set_font('Arial', '', 11)
    pdf.cell(0, 6, 'To Whom It May Concern:', ln=True)
    pdf.ln(6)
    
    intro_text = (
        f"I am writing to request a security freeze be placed on my credit file pursuant to my rights "
        f"under the Fair Credit Reporting Act (FCRA) Section 605A (15 U.S.C. 1681c-1) and applicable "
        f"state law. This freeze will prevent credit reporting agencies from releasing my credit report "
        f"without my authorization."
    )
    pdf.multi_cell(0, 6, intro_text)
    pdf.ln(6)
    
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 8, 'CONSUMER IDENTIFICATION INFORMATION:', ln=True)
    pdf.set_font('Arial', '', 11)
    pdf.ln(2)
    
    pdf.cell(40, 6, 'Full Legal Name:', 0)
    pdf.cell(0, 6, letter_content['client_name'], ln=True)
    
    pdf.cell(40, 6, 'Current Address:', 0)
    pdf.cell(0, 6, letter_content['client_address'], ln=True)
    
    pdf.cell(40, 6, 'Date of Birth:', 0)
    pdf.cell(0, 6, letter_content['dob'], ln=True)
    
    pdf.cell(40, 6, 'SSN (Last 4 digits):', 0)
    pdf.cell(0, 6, f"XXX-XX-{letter_content['ssn_last4']}", ln=True)
    
    pdf.ln(8)
    
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 8, 'LEGAL AUTHORITY:', ln=True)
    pdf.set_font('Arial', '', 11)
    pdf.ln(2)
    
    legal_text = (
        "Under FCRA Section 605A, as amended by the Economic Growth, Regulatory Relief, and Consumer "
        "Protection Act of 2018, consumers have the right to place, temporarily lift, or permanently "
        "remove a security freeze on their credit file at no cost. You are required to place this "
        "freeze within one (1) business day of receiving this request if submitted electronically, "
        "or within three (3) business days if submitted by mail."
    )
    pdf.multi_cell(0, 6, legal_text)
    pdf.ln(6)
    
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 8, 'REQUEST:', ln=True)
    pdf.set_font('Arial', '', 11)
    pdf.ln(2)
    
    request_text = (
        "I hereby request that you:\n\n"
        "1. Place a security freeze on my credit file immediately upon receipt of this letter.\n\n"
        "2. Provide me with written confirmation that the freeze has been placed, including any "
        "PIN, password, or other authentication method I will need to temporarily lift or permanently "
        "remove the freeze in the future.\n\n"
        "3. Send the confirmation to my address listed above within the timeframe required by law."
    )
    pdf.multi_cell(0, 6, request_text)
    pdf.ln(6)
    
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 8, 'ENCLOSURES:', ln=True)
    pdf.set_font('Arial', '', 11)
    pdf.ln(2)
    
    enclosures = [
        "1. Copy of government-issued photo identification (driver's license or state ID)",
        "2. Proof of current address (utility bill, bank statement, or similar document)",
    ]
    for enc in enclosures:
        pdf.cell(0, 6, enc, ln=True)
    
    pdf.ln(8)
    
    notice_text = (
        "Please be advised that failure to comply with this request within the legally mandated "
        "timeframe may constitute a violation of the Fair Credit Reporting Act, subjecting your "
        "organization to statutory damages, actual damages, and attorney's fees."
    )
    pdf.multi_cell(0, 6, notice_text)
    pdf.ln(8)
    
    pdf.cell(0, 6, 'Thank you for your prompt attention to this matter.', ln=True)
    pdf.ln(12)
    
    pdf.cell(0, 6, 'Sincerely,', ln=True)
    pdf.ln(20)
    
    pdf.cell(0, 6, '_' * 40, ln=True)
    pdf.cell(0, 6, letter_content['client_name'], ln=True)
    pdf.cell(0, 6, f"Date: {letter_content['date']}", ln=True)


def _create_freeze_letter_docx(letter_content):
    """
    Create a single freeze letter page for a Word document.
    
    Args:
        letter_content: Dict with letter content from generate_single_freeze_letter
    
    Returns:
        List of paragraph/content data to add to the document
    """
    content_items = []
    
    content_items.append({'type': 'paragraph', 'text': letter_content['date'], 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': letter_content['bureau_name'], 'bold': True, 'size': 11})
    for line in letter_content['bureau_address'].split('\n'):
        if line.strip():
            content_items.append({'type': 'paragraph', 'text': line.strip(), 'bold': False, 'size': 11})
    
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': 'RE: Request for Security Freeze on Credit File', 'bold': True, 'size': 12})
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': 'To Whom It May Concern:', 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    intro_text = (
        f"I am writing to request a security freeze be placed on my credit file pursuant to my rights "
        f"under the Fair Credit Reporting Act (FCRA) Section 605A (15 U.S.C. 1681c-1) and applicable "
        f"state law. This freeze will prevent credit reporting agencies from releasing my credit report "
        f"without my authorization."
    )
    content_items.append({'type': 'paragraph', 'text': intro_text, 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': 'CONSUMER IDENTIFICATION INFORMATION:', 'bold': True, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': f"Full Legal Name: {letter_content['client_name']}", 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': f"Current Address: {letter_content['client_address']}", 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': f"Date of Birth: {letter_content['dob']}", 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': f"SSN (Last 4 digits): XXX-XX-{letter_content['ssn_last4']}", 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': 'LEGAL AUTHORITY:', 'bold': True, 'size': 11})
    legal_text = (
        "Under FCRA Section 605A, as amended by the Economic Growth, Regulatory Relief, and Consumer "
        "Protection Act of 2018, consumers have the right to place, temporarily lift, or permanently "
        "remove a security freeze on their credit file at no cost. You are required to place this "
        "freeze within one (1) business day of receiving this request if submitted electronically, "
        "or within three (3) business days if submitted by mail."
    )
    content_items.append({'type': 'paragraph', 'text': legal_text, 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': 'REQUEST:', 'bold': True, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': "I hereby request that you:", 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': "1. Place a security freeze on my credit file immediately upon receipt of this letter.", 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': "2. Provide me with written confirmation that the freeze has been placed, including any PIN, password, or other authentication method I will need to temporarily lift or permanently remove the freeze in the future.", 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': "3. Send the confirmation to my address listed above within the timeframe required by law.", 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': 'ENCLOSURES:', 'bold': True, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': "1. Copy of government-issued photo identification (driver's license or state ID)", 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': "2. Proof of current address (utility bill, bank statement, or similar document)", 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    notice_text = (
        "Please be advised that failure to comply with this request within the legally mandated "
        "timeframe may constitute a violation of the Fair Credit Reporting Act, subjecting your "
        "organization to statutory damages, actual damages, and attorney's fees."
    )
    content_items.append({'type': 'paragraph', 'text': notice_text, 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': 'Thank you for your prompt attention to this matter.', 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    content_items.append({'type': 'paragraph', 'text': 'Sincerely,', 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    content_items.append({'type': 'spacing'})
    content_items.append({'type': 'paragraph', 'text': '_' * 40, 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': letter_content['client_name'], 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': f"Date: {letter_content['date']}", 'bold': False, 'size': 11})
    
    return content_items


def _add_freeze_letter_to_docx(doc, letter_content, add_page_break=True):
    """
    Add a single freeze letter to the Word document.
    
    Args:
        doc: Document object
        letter_content: Dict with letter content
        add_page_break: Whether to add a page break after the letter
    """
    content_items = _create_freeze_letter_docx(letter_content)
    
    for item in content_items:
        if item['type'] == 'spacing':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
        elif item['type'] == 'paragraph':
            p = doc.add_paragraph()
            run = p.add_run(item['text'])
            run.font.name = 'Arial'
            run.font.size = Pt(item.get('size', 11))
            run.bold = item.get('bold', False)
            p.paragraph_format.space_after = Pt(2)
    
    if add_page_break:
        doc.add_page_break()


def generate_freeze_letters(client_id, bureaus=None):
    """
    Generate bulk freeze letters for specified bureaus (or all 12 if bureaus is None).
    
    Args:
        client_id: ID of the client in the database
        bureaus: List of bureau names to generate letters for, or None for all 12
    
    Returns:
        dict: {
            'success': bool,
            'batch_id': str (UUID),
            'pdf_path': str,
            'bureaus_included': list,
            'error': str (if failed)
        }
    """
    db = get_db()
    
    try:
        client = db.query(Client).filter(Client.id == client_id).first()
        
        if not client:
            return {
                'success': False,
                'error': f'Client with ID {client_id} not found'
            }
        
        if bureaus is None:
            target_bureaus = list(CRA_ADDRESSES.keys())
        else:
            target_bureaus = [b for b in bureaus if b in CRA_ADDRESSES]
            if not target_bureaus:
                return {
                    'success': False,
                    'error': 'No valid bureaus specified'
                }
        
        output_dir = f"static/client_uploads/{client_id}/freeze_letters"
        os.makedirs(output_dir, exist_ok=True)
        
        batch_uuid = str(uuid.uuid4())
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        client_name_safe = ''.join(c for c in (client.name or 'client') if c.isalnum() or c in ' _-').replace(' ', '_')
        pdf_filename = f"{client_name_safe}_Freeze_Letters_{timestamp}.pdf"
        pdf_path = os.path.join(output_dir, pdf_filename)
        
        pdf = FreezeLetterPDF()
        doc = Document()
        
        letter_contents = []
        for bureau_name in target_bureaus:
            bureau_info = CRA_ADDRESSES[bureau_name]
            letter_content = generate_single_freeze_letter(client, bureau_name, bureau_info)
            letter_contents.append(letter_content)
            _add_letter_to_pdf(pdf, letter_content)
        
        pdf.output(pdf_path)
        
        for i, letter_content in enumerate(letter_contents):
            is_last = (i == len(letter_contents) - 1)
            _add_freeze_letter_to_docx(doc, letter_content, add_page_break=not is_last)
        
        docx_filename = f"{client_name_safe}_Freeze_Letters_{timestamp}.docx"
        docx_path = os.path.join(output_dir, docx_filename)
        doc.save(docx_path)
        
        freeze_batch = FreezeLetterBatch(
            client_id=client_id,
            batch_uuid=batch_uuid,
            bureaus_included=target_bureaus,
            total_bureaus=len(target_bureaus),
            generated_pdf_path=pdf_path,
            generated_docx_path=docx_path,
            status='generated',
            created_at=datetime.utcnow()
        )
        
        db.add(freeze_batch)
        db.commit()
        db.refresh(freeze_batch)
        
        return {
            'success': True,
            'batch_id': batch_uuid,
            'batch_db_id': freeze_batch.id,
            'pdf_path': pdf_path,
            'docx_path': docx_path,
            'bureaus_included': target_bureaus,
            'total_letters': len(target_bureaus)
        }
        
    except Exception as e:
        db.rollback()
        return {
            'success': False,
            'error': str(e)
        }
    finally:
        db.close()


def generate_freeze_letters_for_primary_bureaus(client_id):
    """Convenience function to generate freeze letters for only the 3 primary CRAs"""
    primary = list(get_primary_bureaus().keys())
    return generate_freeze_letters(client_id, bureaus=primary)


def generate_freeze_letters_for_secondary_bureaus(client_id):
    """Convenience function to generate freeze letters for only the 9 secondary bureaus"""
    secondary = list(get_secondary_bureaus().keys())
    return generate_freeze_letters(client_id, bureaus=secondary)


def get_freeze_batch_status(batch_uuid):
    """
    Get the status of a freeze letter batch by UUID.
    
    Args:
        batch_uuid: UUID of the batch
    
    Returns:
        dict: Batch information or error
    """
    db = get_db()
    try:
        batch = db.query(FreezeLetterBatch).filter(FreezeLetterBatch.batch_uuid == batch_uuid).first()
        if not batch:
            return {'success': False, 'error': 'Batch not found'}
        
        return {
            'success': True,
            'batch_id': batch.batch_uuid,
            'client_id': batch.client_id,
            'bureaus_included': batch.bureaus_included,
            'total_bureaus': batch.total_bureaus,
            'pdf_path': batch.generated_pdf_path,
            'status': batch.status,
            'mail_method': batch.mail_method,
            'mailed_at': batch.mailed_at.isoformat() if batch.mailed_at else None,
            'created_at': batch.created_at.isoformat() if batch.created_at else None
        }
    finally:
        db.close()


def list_client_freeze_batches(client_id):
    """
    List all freeze letter batches for a specific client.
    
    Args:
        client_id: ID of the client
    
    Returns:
        list: List of batch summaries
    """
    db = get_db()
    try:
        batches = db.query(FreezeLetterBatch).filter(
            FreezeLetterBatch.client_id == client_id
        ).order_by(FreezeLetterBatch.created_at.desc()).all()
        
        return [{
            'batch_id': b.batch_uuid,
            'total_bureaus': b.total_bureaus,
            'pdf_path': b.generated_pdf_path,
            'status': b.status,
            'created_at': b.created_at.isoformat() if b.created_at else None
        } for b in batches]
    finally:
        db.close()
