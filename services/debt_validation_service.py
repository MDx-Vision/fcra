"""
Debt Validation Letter Service
Generates FDCPA-compliant debt validation letters for collection accounts.
Auto-populates client PII and can be triggered automatically or on-demand.
Uses fpdf2 for PDF generation and python-docx for Word document generation.
"""
import os
import uuid
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from database import get_db, Client, Case, Violation

try:
    from fpdf.pattern import LinearGradient
    HAS_GRADIENTS = True
except ImportError:
    HAS_GRADIENTS = False

TEAL = (0, 128, 128)
TEAL_HEX = "#008080"
LIME_HEX = "#32CD32"
WHITE = (255, 255, 255)
DARK_GRAY = (51, 51, 51)
MEDIUM_GRAY = (102, 102, 102)


COLLECTION_AGENCY_TEMPLATE = {
    'generic': {
        'name': '[COLLECTION AGENCY NAME]',
        'address': '[COLLECTION AGENCY ADDRESS]',
        'city_state_zip': '[CITY, STATE ZIP]'
    }
}


class DebtValidationPDF(FPDF):
    """Custom PDF class for debt validation letters with Brightpath Ascend branding"""
    
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=25)
        self.set_margins(left=20, top=20, right=20)
    
    def header(self):
        header_height = 30
        
        if HAS_GRADIENTS:
            try:
                gradient = LinearGradient(
                    self,
                    from_x=0, from_y=0,
                    to_x=self.w, to_y=0,
                    colors=[TEAL_HEX, LIME_HEX]
                )
                with self.use_pattern(gradient):
                    self.rect(0, 0, self.w, header_height, style="F")
            except Exception:
                self.set_fill_color(*TEAL)
                self.rect(0, 0, self.w, header_height, style="F")
        else:
            self.set_fill_color(*TEAL)
            self.rect(0, 0, self.w, header_height, style="F")
        
        self.set_y(7)
        self.set_font('Helvetica', 'B', 16)
        self.set_text_color(*WHITE)
        self.cell(0, 8, "BRIGHTPATH ASCEND GROUP", align='C', new_x="LMARGIN", new_y="NEXT")
        
        self.set_font('Helvetica', '', 8)
        self.cell(0, 5, "Consumer Protection & FCRA Litigation Specialists", align='C')
        
        self.set_text_color(0, 0, 0)
        self.set_y(header_height + 5)
    
    def footer(self):
        self.set_y(-18)
        
        self.set_draw_color(*TEAL)
        self.set_line_width(0.5)
        self.line(20, self.get_y(), self.w - 20, self.get_y())
        
        self.set_y(-13)
        self.set_font('Helvetica', '', 8)
        self.set_text_color(*MEDIUM_GRAY)
        self.cell(0, 5, f'Page {self.page_no()}', align='C', new_x="LMARGIN", new_y="NEXT")
        self.set_font('Helvetica', 'I', 7)
        self.cell(0, 4, "Brightpath Ascend Group | www.brightpathascend.com", align='C')
        
        self.set_text_color(0, 0, 0)
        self.set_line_width(0.2)


def get_client_pii(client):
    """Extract PII from client object or dict"""
    if isinstance(client, dict):
        return {
            'name': client.get('name', '') or f"{client.get('first_name', '')} {client.get('last_name', '')}".strip(),
            'first_name': client.get('first_name', ''),
            'last_name': client.get('last_name', ''),
            'address_street': client.get('address_street', ''),
            'address_city': client.get('address_city', ''),
            'address_state': client.get('address_state', ''),
            'address_zip': client.get('address_zip', ''),
            'ssn_last_four': client.get('ssn_last_four', 'XXXX'),
            'date_of_birth': client.get('date_of_birth', None),
            'email': client.get('email', ''),
            'phone': client.get('phone', '')
        }
    else:
        name = client.name or ''
        if not name and client.first_name and client.last_name:
            name = f"{client.first_name} {client.last_name}"
        return {
            'name': name,
            'first_name': client.first_name or '',
            'last_name': client.last_name or '',
            'address_street': client.address_street or '',
            'address_city': client.address_city or '',
            'address_state': client.address_state or '',
            'address_zip': client.address_zip or '',
            'ssn_last_four': client.ssn_last_four or 'XXXX',
            'date_of_birth': client.date_of_birth,
            'email': client.email or '',
            'phone': client.phone or ''
        }


def generate_debt_validation_letter(client, collection_info):
    """
    Generate a single debt validation letter for a collection account.
    
    Args:
        client: Client object or dict with client information
        collection_info: Dict with collection account details:
            - creditor_name: Name of collection agency
            - creditor_address: Address of collection agency (optional)
            - original_creditor: Original creditor name
            - account_number: Account number (will be partially masked)
            - balance: Current balance claimed
            - date_opened: Date account was opened/placed
    
    Returns:
        dict: Letter content for PDF generation
    """
    pii = get_client_pii(client)
    
    if pii['date_of_birth']:
        if hasattr(pii['date_of_birth'], 'strftime'):
            dob_formatted = pii['date_of_birth'].strftime('%B %d, %Y')
        else:
            dob_formatted = str(pii['date_of_birth'])
    else:
        dob_formatted = '[DATE OF BIRTH]'
    
    client_full_address = f"{pii['address_street']}"
    client_city_state_zip = f"{pii['address_city']}, {pii['address_state']} {pii['address_zip']}"
    
    if not pii['address_street']:
        client_full_address = '[CLIENT ADDRESS]'
        client_city_state_zip = '[CITY, STATE ZIP]'
    
    creditor_name = collection_info.get('creditor_name', '[COLLECTION AGENCY]')
    creditor_address = collection_info.get('creditor_address', '[ADDRESS]')
    creditor_city_state_zip = collection_info.get('creditor_city_state_zip', '[CITY, STATE ZIP]')
    original_creditor = collection_info.get('original_creditor', '[ORIGINAL CREDITOR]')
    account_number = collection_info.get('account_number', '[ACCOUNT NUMBER]')
    balance = collection_info.get('balance', '[AMOUNT]')
    
    if account_number and len(str(account_number)) > 4:
        masked_account = f"XXXX{str(account_number)[-4:]}"
    else:
        masked_account = str(account_number)
    
    if isinstance(balance, (int, float)):
        balance_formatted = f"${balance:,.2f}"
    else:
        balance_formatted = str(balance)
    
    today = datetime.now().strftime('%B %d, %Y')
    
    return {
        'date': today,
        'creditor_name': creditor_name,
        'creditor_address': creditor_address,
        'creditor_city_state_zip': creditor_city_state_zip,
        'client_name': pii['name'],
        'client_address': client_full_address,
        'client_city_state_zip': client_city_state_zip,
        'original_creditor': original_creditor,
        'account_number': masked_account,
        'balance': balance_formatted,
        'ssn_last4': pii['ssn_last_four'],
        'reference_number': collection_info.get('reference_number', '')
    }


def _add_validation_letter_to_pdf(pdf, letter_content):
    """Add a single debt validation letter to the PDF document"""
    pdf.add_page()
    
    pdf.set_font('Arial', '', 11)
    pdf.cell(0, 6, letter_content['client_name'], ln=True)
    pdf.cell(0, 6, letter_content['client_address'], ln=True)
    pdf.cell(0, 6, letter_content['client_city_state_zip'], ln=True)
    pdf.ln(6)
    pdf.cell(0, 6, letter_content['date'], ln=True)
    pdf.ln(8)
    
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 6, letter_content['creditor_name'], ln=True)
    pdf.set_font('Arial', '', 11)
    pdf.cell(0, 6, letter_content['creditor_address'], ln=True)
    pdf.cell(0, 6, letter_content['creditor_city_state_zip'], ln=True)
    pdf.ln(8)
    
    if letter_content.get('reference_number'):
        pdf.cell(0, 6, f"RE: Account Reference #{letter_content['reference_number']}", ln=True)
    pdf.cell(0, 6, f"RE: Alleged Account #{letter_content['account_number']}", ln=True)
    pdf.cell(0, 6, f"Alleged Original Creditor: {letter_content['original_creditor']}", ln=True)
    pdf.cell(0, 6, f"Alleged Balance: {letter_content['balance']}", ln=True)
    pdf.ln(8)
    
    pdf.set_font('Arial', 'B', 12)
    pdf.set_text_color(139, 0, 0)
    pdf.cell(0, 8, 'NOTICE OF DISPUTE AND DEMAND FOR DEBT VALIDATION', ln=True, align='C')
    pdf.cell(0, 6, 'PURSUANT TO 15 U.S.C. 1692g (FDCPA)', ln=True, align='C')
    pdf.set_text_color(0, 0, 0)
    pdf.ln(8)
    
    pdf.set_font('Arial', '', 11)
    pdf.cell(0, 6, 'To Whom It May Concern:', ln=True)
    pdf.ln(4)
    
    dispute_text = (
        "I am writing in response to your communication regarding the above-referenced account. "
        "I DISPUTE THIS ALLEGED DEBT IN ITS ENTIRETY. This letter serves as my formal written dispute "
        "and demand for validation of this alleged debt pursuant to the Fair Debt Collection Practices "
        "Act, 15 U.S.C. 1692g."
    )
    pdf.multi_cell(0, 6, dispute_text)
    pdf.ln(4)
    
    cease_text = (
        "PLEASE BE ADVISED: Until you provide the validation requested below, you must CEASE AND DESIST "
        "all collection activities. Any continued collection attempts prior to validation will constitute "
        "a violation of the FDCPA, subjecting your company to statutory damages of up to $1,000 per "
        "violation, plus actual damages and attorney's fees."
    )
    pdf.multi_cell(0, 6, cease_text)
    pdf.ln(6)
    
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 8, 'REQUIRED DOCUMENTATION FOR PROPER VALIDATION:', ln=True)
    pdf.set_font('Arial', '', 11)
    pdf.ln(2)
    
    validation_items = [
        "1. Complete payment history from the original creditor showing all charges, payments, and fees",
        "2. A copy of the original signed contract or agreement bearing my signature",
        "3. Documentation proving the chain of assignment/ownership from original creditor to you",
        "4. Verification that the statute of limitations has not expired on this alleged debt",
        "5. Documentation showing your license to collect debts in my state of residence",
        "6. Proof that you have authority to collect this specific debt",
        "7. Name and address of the original creditor",
        "8. The exact amount of the alleged debt broken down by principal, interest, and fees",
        "9. Documentation showing the date of first delinquency (DOFD)",
        "10. If applicable, assignment or bill of sale from the original creditor to your company"
    ]
    
    for item in validation_items:
        pdf.multi_cell(0, 5, item)
        pdf.ln(1)
    
    pdf.ln(4)
    
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 8, 'ADDITIONAL DEMANDS:', ln=True)
    pdf.set_font('Arial', '', 11)
    pdf.ln(2)
    
    demands = [
        "- CEASE all telephone calls to me, my family, and my place of employment",
        "- DO NOT report this disputed debt to any credit reporting agency",
        "- If already reported, REMOVE this tradeline from all credit reports immediately",
        "- Communicate with me ONLY in writing at the address listed above"
    ]
    
    for demand in demands:
        pdf.cell(0, 6, demand, ln=True)
    
    pdf.ln(6)
    
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 8, 'LEGAL NOTICE:', ln=True)
    pdf.set_font('Arial', '', 11)
    pdf.ln(2)
    
    legal_notice = (
        "Be advised that I am aware of my rights under the Fair Debt Collection Practices Act (FDCPA), "
        "the Fair Credit Reporting Act (FCRA), and applicable state laws. Any violation of these laws "
        "will be documented and may be used in legal proceedings against your company. "
        "This letter is dated and will be sent via certified mail to preserve evidence of your receipt."
    )
    pdf.multi_cell(0, 6, legal_notice)
    pdf.ln(4)
    
    response_text = (
        "You have 30 days from receipt of this letter to provide the requested validation. Failure to "
        "respond within this timeframe shall be deemed an acknowledgment that this debt is invalid and "
        "must be deleted from all credit reporting agencies."
    )
    pdf.multi_cell(0, 6, response_text)
    pdf.ln(8)
    
    pdf.cell(0, 6, 'Govern yourself accordingly.', ln=True)
    pdf.ln(12)
    
    pdf.cell(0, 6, 'Sincerely,', ln=True)
    pdf.ln(20)
    
    pdf.cell(0, 6, '_' * 40, ln=True)
    pdf.cell(0, 6, letter_content['client_name'], ln=True)
    pdf.ln(10)
    
    pdf.set_font('Arial', 'I', 9)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 5, 'CC: Consumer Financial Protection Bureau', ln=True)
    pdf.cell(0, 5, '    State Attorney General Office', ln=True)
    pdf.cell(0, 5, '    Federal Trade Commission', ln=True)


def _create_validation_letter_docx(letter_content):
    """
    Create content for a single debt validation letter for a Word document.
    
    Args:
        letter_content: Dict with letter content from generate_debt_validation_letter
    
    Returns:
        List of paragraph/content data to add to the document
    """
    content_items = []
    
    content_items.append({'type': 'paragraph', 'text': letter_content['client_name'], 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': letter_content['client_address'], 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': letter_content['client_city_state_zip'], 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    content_items.append({'type': 'paragraph', 'text': letter_content['date'], 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': letter_content['creditor_name'], 'bold': True, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': letter_content['creditor_address'], 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': letter_content['creditor_city_state_zip'], 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    if letter_content.get('reference_number'):
        content_items.append({'type': 'paragraph', 'text': f"RE: Account Reference #{letter_content['reference_number']}", 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': f"RE: Alleged Account #{letter_content['account_number']}", 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': f"Alleged Original Creditor: {letter_content['original_creditor']}", 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': f"Alleged Balance: {letter_content['balance']}", 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': 'NOTICE OF DISPUTE AND DEMAND FOR DEBT VALIDATION', 'bold': True, 'size': 12, 'center': True})
    content_items.append({'type': 'paragraph', 'text': 'PURSUANT TO 15 U.S.C. 1692g (FDCPA)', 'bold': True, 'size': 12, 'center': True})
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': 'To Whom It May Concern:', 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    dispute_text = (
        "I am writing in response to your communication regarding the above-referenced account. "
        "I DISPUTE THIS ALLEGED DEBT IN ITS ENTIRETY. This letter serves as my formal written dispute "
        "and demand for validation of this alleged debt pursuant to the Fair Debt Collection Practices "
        "Act, 15 U.S.C. 1692g."
    )
    content_items.append({'type': 'paragraph', 'text': dispute_text, 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    cease_text = (
        "PLEASE BE ADVISED: Until you provide the validation requested below, you must CEASE AND DESIST "
        "all collection activities. Any continued collection attempts prior to validation will constitute "
        "a violation of the FDCPA, subjecting your company to statutory damages of up to $1,000 per "
        "violation, plus actual damages and attorney's fees."
    )
    content_items.append({'type': 'paragraph', 'text': cease_text, 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': 'REQUIRED DOCUMENTATION FOR PROPER VALIDATION:', 'bold': True, 'size': 11})
    
    validation_items = [
        "1. Complete payment history from the original creditor showing all charges, payments, and fees",
        "2. A copy of the original signed contract or agreement bearing my signature",
        "3. Documentation proving the chain of assignment/ownership from original creditor to you",
        "4. Verification that the statute of limitations has not expired on this alleged debt",
        "5. Documentation showing your license to collect debts in my state of residence",
        "6. Proof that you have authority to collect this specific debt",
        "7. Name and address of the original creditor",
        "8. The exact amount of the alleged debt broken down by principal, interest, and fees",
        "9. Documentation showing the date of first delinquency (DOFD)",
        "10. If applicable, assignment or bill of sale from the original creditor to your company"
    ]
    for item in validation_items:
        content_items.append({'type': 'paragraph', 'text': item, 'bold': False, 'size': 11})
    
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': 'ADDITIONAL DEMANDS:', 'bold': True, 'size': 11})
    demands = [
        "- CEASE all telephone calls to me, my family, and my place of employment",
        "- DO NOT report this disputed debt to any credit reporting agency",
        "- If already reported, REMOVE this tradeline from all credit reports immediately",
        "- Communicate with me ONLY in writing at the address listed above"
    ]
    for demand in demands:
        content_items.append({'type': 'paragraph', 'text': demand, 'bold': False, 'size': 11})
    
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': 'LEGAL NOTICE:', 'bold': True, 'size': 11})
    legal_notice = (
        "Be advised that I am aware of my rights under the Fair Debt Collection Practices Act (FDCPA), "
        "the Fair Credit Reporting Act (FCRA), and applicable state laws. Any violation of these laws "
        "will be documented and may be used in legal proceedings against your company. "
        "This letter is dated and will be sent via certified mail to preserve evidence of your receipt."
    )
    content_items.append({'type': 'paragraph', 'text': legal_notice, 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    response_text = (
        "You have 30 days from receipt of this letter to provide the requested validation. Failure to "
        "respond within this timeframe shall be deemed an acknowledgment that this debt is invalid and "
        "must be deleted from all credit reporting agencies."
    )
    content_items.append({'type': 'paragraph', 'text': response_text, 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': 'Govern yourself accordingly.', 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    content_items.append({'type': 'paragraph', 'text': 'Sincerely,', 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    content_items.append({'type': 'spacing'})
    content_items.append({'type': 'paragraph', 'text': '_' * 40, 'bold': False, 'size': 11})
    content_items.append({'type': 'paragraph', 'text': letter_content['client_name'], 'bold': False, 'size': 11})
    content_items.append({'type': 'spacing'})
    
    content_items.append({'type': 'paragraph', 'text': 'CC: Consumer Financial Protection Bureau', 'bold': False, 'size': 9, 'italic': True})
    content_items.append({'type': 'paragraph', 'text': '    State Attorney General Office', 'bold': False, 'size': 9, 'italic': True})
    content_items.append({'type': 'paragraph', 'text': '    Federal Trade Commission', 'bold': False, 'size': 9, 'italic': True})
    
    return content_items


def _add_validation_letter_to_docx(doc, letter_content, add_page_break=True):
    """
    Add a single debt validation letter to the Word document.
    
    Args:
        doc: Document object
        letter_content: Dict with letter content
        add_page_break: Whether to add a page break after the letter
    """
    content_items = _create_validation_letter_docx(letter_content)
    
    for item in content_items:
        if item['type'] == 'spacing':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
        elif item['type'] == 'paragraph':
            p = doc.add_paragraph()
            run = p.add_run(item['text'])
            run.font.name = 'Arial'
            run.font.size = Pt(item.get('size', 11))
            run.bold = item.get('bold', False)
            run.italic = item.get('italic', False)
            p.paragraph_format.space_after = Pt(2)
            if item.get('center'):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if add_page_break:
        doc.add_page_break()


def generate_validation_letters(client_id, collections=None, case_id=None):
    """
    Generate debt validation letters for collection accounts.
    
    Args:
        client_id: ID of the client in the database
        collections: List of collection dicts, or None to auto-detect from case violations
        case_id: Case ID to pull collection violations from (if collections not provided)
    
    Returns:
        dict: {
            'success': bool,
            'pdf_path': str,
            'letters_generated': int,
            'collections': list,
            'error': str (if failed)
        }
    """
    db = get_db()
    
    try:
        client = db.query(Client).filter(Client.id == client_id).first()
        
        if not client:
            return {
                'success': False,
                'error': f'Client with ID {client_id} not found'
            }
        
        if collections is None and case_id:
            case = db.query(Case).filter(Case.id == case_id).first()
            if case:
                violations = db.query(Violation).filter(
                    Violation.case_id == case_id
                ).all()
                
                collections = []
                seen_accounts = set()
                
                for v in violations:
                    if v.creditor_name and v.creditor_name not in seen_accounts:
                        if any(kw in (v.violation_type or '').lower() for kw in ['collection', 'debt', 'charged']):
                            collections.append({
                                'creditor_name': v.creditor_name,
                                'creditor_address': '',
                                'creditor_city_state_zip': '',
                                'original_creditor': v.creditor_name,
                                'account_number': v.account_number or '',
                                'balance': 0
                            })
                            seen_accounts.add(v.creditor_name)
        
        if not collections:
            return {
                'success': False,
                'error': 'No collection accounts provided or found'
            }
        
        output_dir = f"static/client_uploads/{client_id}/validation_letters"
        os.makedirs(output_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        client_name_safe = ''.join(c for c in (client.name or 'client') if c.isalnum() or c in ' _-').replace(' ', '_')
        pdf_filename = f"{client_name_safe}_Debt_Validation_Letters_{timestamp}.pdf"
        pdf_path = os.path.join(output_dir, pdf_filename)
        
        pdf = DebtValidationPDF()
        doc = Document()
        
        letter_contents = []
        for collection in collections:
            letter_content = generate_debt_validation_letter(client, collection)
            letter_contents.append(letter_content)
            _add_validation_letter_to_pdf(pdf, letter_content)
        
        pdf.output(pdf_path)
        
        for i, letter_content in enumerate(letter_contents):
            is_last = (i == len(letter_contents) - 1)
            _add_validation_letter_to_docx(doc, letter_content, add_page_break=not is_last)
        
        docx_filename = f"{client_name_safe}_Debt_Validation_Letters_{timestamp}.docx"
        docx_path = os.path.join(output_dir, docx_filename)
        doc.save(docx_path)
        
        return {
            'success': True,
            'pdf_path': pdf_path,
            'docx_path': docx_path,
            'letters_generated': len(collections),
            'collections': [c.get('creditor_name', 'Unknown') for c in collections]
        }
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'error': str(e)
        }
    finally:
        db.close()


def generate_validation_letter_single(client_id, collection_info):
    """
    Generate a single debt validation letter for one collection account.
    
    Args:
        client_id: ID of the client
        collection_info: Dict with collection account details
    
    Returns:
        dict: Result with pdf_path or error
    """
    return generate_validation_letters(client_id, collections=[collection_info])


def auto_generate_validation_letters_from_analysis(case_id):
    """
    Automatically generate validation letters for all collection accounts
    detected during credit report analysis.
    
    Args:
        case_id: ID of the case with completed analysis
    
    Returns:
        dict: Result with pdf_path and collection count
    """
    db = get_db()
    
    try:
        case = db.query(Case).filter(Case.id == case_id).first()
        if not case:
            return {'success': False, 'error': 'Case not found'}
        
        client = db.query(Client).filter(Client.id == case.client_id).first()
        if not client:
            return {'success': False, 'error': 'Client not found'}
        
        violations = db.query(Violation).filter(Violation.case_id == case_id).all()
        
        collections = []
        seen_accounts = set()
        
        collection_keywords = ['collection', 'collections', 'debt collector', 'charged off', 
                              'charge-off', 'sold', 'purchased debt', 'assigned']
        
        for v in violations:
            if v.creditor_name and v.creditor_name.lower() not in seen_accounts:
                is_collection = False
                
                desc_lower = (v.description or '').lower()
                type_lower = (v.violation_type or '').lower()
                
                for kw in collection_keywords:
                    if kw in desc_lower or kw in type_lower:
                        is_collection = True
                        break
                
                if is_collection:
                    collections.append({
                        'creditor_name': v.creditor_name,
                        'creditor_address': '',
                        'creditor_city_state_zip': '',
                        'original_creditor': v.creditor_name,
                        'account_number': v.account_number or '',
                        'balance': 0
                    })
                    seen_accounts.add(v.creditor_name.lower())
        
        if not collections:
            return {
                'success': True,
                'message': 'No collection accounts detected in analysis',
                'letters_generated': 0
            }
        
        return generate_validation_letters(case.client_id, collections=collections)
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}
    finally:
        db.close()


def get_common_collection_agencies():
    """Return list of common collection agencies with addresses for quick selection"""
    return [
        {
            'name': 'Portfolio Recovery Associates',
            'address': '120 Corporate Blvd.',
            'city_state_zip': 'Norfolk, VA 23502'
        },
        {
            'name': 'Midland Credit Management',
            'address': '350 Camino de la Reina, Suite 100',
            'city_state_zip': 'San Diego, CA 92108'
        },
        {
            'name': 'LVNV Funding LLC',
            'address': '15 South 700 East, Suite 200',
            'city_state_zip': 'Las Vegas, NV 89101'
        },
        {
            'name': 'Cavalry SPV I, LLC',
            'address': '500 Summit Lake Drive',
            'city_state_zip': 'Valhalla, NY 10595'
        },
        {
            'name': 'Enhanced Recovery Company (ERC)',
            'address': '8014 Bayberry Road',
            'city_state_zip': 'Jacksonville, FL 32256'
        },
        {
            'name': 'IC System',
            'address': '444 Highway 96 E',
            'city_state_zip': 'Saint Paul, MN 55164'
        },
        {
            'name': 'Transworld Systems Inc.',
            'address': '2 Crossways Park Dr N',
            'city_state_zip': 'Woodbury, NY 11797'
        },
        {
            'name': 'Convergent Outsourcing',
            'address': '5200 Regal Dr, Suite 100',
            'city_state_zip': 'Lake Oswego, OR 97035'
        }
    ]
